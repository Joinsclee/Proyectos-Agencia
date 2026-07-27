from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import build_matriz_hu_cliente as ui


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "output" / "docs" / "Panorama_General_Historias_Usuario_Radar_2026-07-26.docx"


def replace_paragraph_text(paragraph, text, *, size=8.5, color=ui.MUTED, bold=False, align=None):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    if align is not None:
        paragraph.alignment = align
    ui.set_font(paragraph.add_run(text), size=size, color=color, bold=bold)


def configure_client_document(doc):
    # Preset: standard_business_brief.
    # Named visual override: Radar CRECE purple/gold brand palette.
    # First-page header pattern: customer_pack.
    ui.configure_document(doc)
    section = doc.sections[0]

    header_table = section.header.tables[0]
    replace_paragraph_text(
        header_table.cell(0, 0).paragraphs[0],
        "RADAR CRECE  |  HISTORIAS DE USUARIO",
        bold=True,
    )
    replace_paragraph_text(
        header_table.cell(0, 1).paragraphs[0],
        "Panorama funcional",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )

    footer_table = section.footer.tables[0]
    replace_paragraph_text(
        footer_table.cell(0, 0).paragraphs[0],
        "26 de julio de 2026  •  Documento funcional para conversación con el cliente",
    )

    core = doc.core_properties
    core.title = "Panorama general de historias de usuario — Radar de Oportunidades"
    core.subject = "Catálogo funcional de las historias de usuario de Fases 1 y 2"
    core.author = "Equipo de desarrollo Radar CRECE"
    core.keywords = "Radar, CRECE, historias de usuario, alcance funcional, fase 1, fase 2"
    core.comments = "Documento cualitativo sin porcentajes de avance."


def add_metadata_grid(doc):
    table = doc.add_table(rows=2, cols=2)
    ui.set_table_geometry(table, [4680, 4680], indent=0)
    metadata = [
        ("Preparado para", "Cliente Radar de Oportunidades"),
        ("Propósito", "Alinear alcance y lenguaje funcional"),
        ("Cobertura", "Historias de usuario de Fases 1 y 2"),
        ("Enfoque", "Paneo cualitativo, sin porcentajes"),
    ]
    for idx, (label, value) in enumerate(metadata):
        cell = table.cell(idx // 2, idx % 2)
        ui.set_cell_shading(cell, ui.GRAY)
        ui.set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": ui.MID_GRAY},
            bottom={"val": "single", "sz": "4", "color": ui.MID_GRAY},
            start={"val": "single", "sz": "4", "color": ui.MID_GRAY},
            end={"val": "single", "sz": "4", "color": ui.MID_GRAY},
        )
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = ui.Pt(2)
        ui.set_font(p1.add_run(label.upper()), size=8, color=ui.MUTED, bold=True)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = ui.Pt(0)
        ui.set_font(p2.add_run(value), size=10, color=ui.INK, bold=True)


def add_phase_map_table(doc, rows, header_fill):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [1220, 1740, 3650, 2750]
    ui.set_table_geometry(table, widths)
    headers = ["ID", "Actor principal", "Necesidad", "Resultado esperado"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        ui.set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT
        ui.set_font(p.add_run(text), size=9, color=ui.WHITE, bold=True)
    ui.set_repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = ui.Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT
            ui.set_font(p.add_run(text), size=8.7, bold=(idx == 0))
    return table


def add_acceptance_item(doc, bullet_id, text):
    ui.add_list_item(doc, text, bullet_id)


def add_hu_catalog_block(
    doc,
    bullet_id,
    hu_id,
    title,
    story,
    value,
    result,
    acceptance,
    demonstration,
    decision,
    *,
    phase=1,
):
    ui.add_heading(doc, f"{hu_id} — {title}", 2)
    accent = ui.PURPLE if phase == 1 else ui.BLUE
    fill = ui.PURPLE_LIGHT if phase == 1 else "EEF4FA"
    ui.add_callout(doc, "Historia de usuario", story, fill=fill, accent=accent)
    ui.add_body(doc, f"Valor para el usuario. {value}", bold_prefix="Valor para el usuario. ")
    ui.add_body(doc, f"Resultado esperado. {result}", bold_prefix="Resultado esperado. ")
    ui.add_body(doc, "Criterios de aceptación principales:", bold_prefix="Criterios de aceptación principales:")
    for item in acceptance:
        add_acceptance_item(doc, bullet_id, item)
    ui.add_body(doc, f"Demostración funcional. {demonstration}", bold_prefix="Demostración funcional. ")
    ui.add_callout(
        doc,
        "Decisión o dependencia",
        decision,
        fill=ui.GOLD_LIGHT,
        accent=ui.AMBER,
    )


def add_phase_close(doc, label, text):
    ui.add_callout(doc, label, text, fill=ui.PURPLE_LIGHT, accent=ui.PURPLE)


def main():
    doc = Document()
    configure_client_document(doc)
    bullet_id = ui.add_numbering(doc, "bullet")
    decimal_id = ui.add_numbering(doc, "decimal")
    workflow_decimal_id = ui.add_numbering(doc, "decimal")

    # First-page customer pack.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = ui.Pt(14)
    p.paragraph_format.space_after = ui.Pt(2)
    ui.set_font(p.add_run("DOCUMENTO FUNCIONAL PARA CLIENTE"), size=9.5, color=ui.AMBER, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = ui.Pt(6)
    ui.set_font(
        title.add_run("Panorama general de Historias de Usuario"),
        size=27,
        color=ui.INK,
        bold=True,
    )

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = ui.Pt(18)
    ui.set_font(
        subtitle.add_run("Radar de Oportunidades · Alcance funcional de Fases 1 y 2"),
        size=13.5,
        color=ui.PURPLE,
        bold=True,
    )

    add_metadata_grid(doc)
    doc.add_paragraph()
    ui.add_callout(
        doc,
        "Objetivo del documento",
        "Presentar de manera clara qué necesidades de usuario atiende el producto, cómo se relacionan "
        "las historias entre sí y qué criterios deben validarse con el cliente. Este documento no usa "
        "porcentajes de avance ni pretende sustituir un acta formal de aceptación.",
        fill=ui.PURPLE_LIGHT,
        accent=ui.PURPLE,
    )
    ui.add_body(
        doc,
        "La propuesta central del Radar es reunir oportunidades inmobiliarias dispersas, normalizar sus "
        "datos y convertirlas en información comparable para que el usuario pueda decidir con menos "
        "tiempo, más contexto y menor riesgo.",
    )

    doc.add_page_break()
    ui.add_heading(doc, "1. Cómo se organizan las historias de usuario", 1)
    ui.add_body(
        doc,
        "Una historia de usuario describe una necesidad desde la perspectiva de quien recibe el valor. "
        "Su función es conectar una decisión de negocio con un resultado observable, evitando que el "
        "alcance quede reducido a una lista de tareas técnicas.",
    )
    for text in [
        "Actor: persona o rol que necesita el resultado.",
        "Necesidad: acción que el actor quiere realizar dentro del producto.",
        "Valor: razón por la que esa acción mejora su decisión o experiencia.",
        "Criterios de aceptación: condiciones observables que permiten validar la historia.",
        "Decisión o dependencia: definición del cliente o insumo externo necesario para completar el alcance.",
    ]:
        ui.add_list_item(doc, text, bullet_id)

    ui.add_heading(doc, "1.1 Actores principales", 2)
    for text, prefix in [
        ("Visitante: explora el producto y conoce el valor del Radar antes de registrarse.", "Visitante:"),
        ("Usuario registrado: guarda preferencias, consulta oportunidades y recibe seguimiento.", "Usuario registrado:"),
        ("Inversionista o usuario Pro: accede al análisis completo, comparables y herramientas de decisión.", "Inversionista o usuario Pro:"),
        ("Administrador: controla usuarios, intereses comerciales y operación del servicio.", "Administrador:"),
        ("Operación de datos: mantiene fuentes, calidad, vigencia y trazabilidad de la información.", "Operación de datos:"),
    ]:
        ui.add_list_item(doc, text, bullet_id, bold_prefix=prefix)

    ui.add_heading(doc, "1.2 Recorrido funcional del producto", 2)
    for text in [
        "Descubrir oportunidades provenientes de portales, bancos y remates.",
        "Filtrar por ciudad, zona, tipo, precio y características.",
        "Analizar precio, comparables, riesgo jurídico y rentabilidad.",
        "Guardar criterios y recibir alertas relevantes.",
        "Acceder a beneficios Pro y, cuando aplique, activar un plan.",
        "Administrar la operación, permisos y calidad del servicio.",
    ]:
        ui.add_list_item(doc, text, decimal_id)

    doc.add_page_break()
    ui.add_heading(doc, "2. Mapa general de historias por fase", 1)
    ui.add_body(
        doc,
        "La Fase 1 consolida el motor de información y análisis. La Fase 2 convierte ese motor en un "
        "producto accesible, operable y comercializable. Las dos fases forman una sola experiencia.",
    )

    ui.add_heading(doc, "2.1 Fase 1 — Datos, comparables y riesgo", 2)
    add_phase_map_table(
        doc,
        [
            ("HU-F1-01", "Inversionista", "Comparar una propiedad con mercado equivalente.", "Clasificación CRECE con soporte estadístico."),
            ("HU-F1-02", "Usuario", "Conocer vigencia y cambios de publicaciones.", "Inventario confiable y trazable."),
            ("HU-F1-03", "Inversionista", "Analizar activos bancarios con datos incompletos.", "Comparación útil sin descartar oportunidades."),
            ("HU-F1-04", "Inversionista", "Consultar activos AVAL desde boletines PDF.", "Fichas normalizadas y cambios identificables."),
            ("HU-F1-05", "Inversionista", "Entender riesgo y condiciones de un remate.", "Lectura jurídica y comercial más segura."),
        ],
        ui.INK,
    )

    ui.add_heading(doc, "2.2 Fase 2 — Acceso, producto y operación", 2)
    add_phase_map_table(
        doc,
        [
            ("HU-F2-01", "Usuario", "Ingresar con Google.", "Acceso simple y sesión persistente."),
            ("HU-F2-02", "Usuario / Pro", "Entender y usar beneficios por plan.", "Experiencia freemium y Pro coherente."),
            ("HU-F2-03", "Usuario", "Recibir alertas por correo.", "Reacción oportuna ante nuevas oportunidades."),
            ("HU-F2-04", "Administrador", "Controlar usuarios y operación.", "Gobierno interno y trazabilidad."),
            ("HU-F2-05", "Inversionista", "Estimar canon y rentabilidad.", "Análisis integral de inversión."),
            ("HU-F2-06", "Usuario / Pro", "Pagar y activar un plan.", "Suscripción segura y conciliable."),
            ("HU-F2-07", "Visitante / Usuario", "Filtrar con claridad en desktop y móvil.", "Búsqueda eficiente y comprensible."),
        ],
        ui.PURPLE,
    )

    doc.add_page_break()
    ui.add_heading(doc, "3. Catálogo funcional — Fase 1", 1)
    ui.add_body(
        doc,
        "Las historias de Fase 1 definen la confianza del producto: calidad de las fuentes, comparabilidad "
        "del mercado y lectura del riesgo. Si esta base no es consistente, las funciones comerciales de "
        "Fase 2 pierden valor.",
    )

    add_hu_catalog_block(
        doc,
        bullet_id,
        "HU-F1-01",
        "Núcleo de comparables e Índice CRECE",
        "Como inversionista inmobiliario, quiero comparar una propiedad con inmuebles equivalentes del "
        "mercado abierto para saber si su precio representa una oportunidad real.",
        "Reduce la dependencia de una lectura subjetiva del precio y permite priorizar oportunidades.",
        "Una propiedad se compara contra una muestra pertinente y recibe una lectura CRECE comprensible.",
        [
            "Usar precio por metro cuadrado y una mediana robusta.",
            "Recortar valores atípicos y comparar por tipo, área, atributos y cercanía.",
            "Excluir bancos y remates del universo de referencia del mercado abierto.",
            "Explicar la clasificación obtenida con criterios visibles.",
        ],
        "Seleccionar una propiedad y revisar su muestra de comparables, mediana y clasificación CRECE.",
        "Aprobar la regla de ampliación geográfica cuando no exista muestra suficiente en el barrio.",
        phase=1,
    )
    doc.add_page_break()
    add_hu_catalog_block(
        doc,
        bullet_id,
        "HU-F1-02",
        "Rango de precio y vigencia del Portal",
        "Como usuario del Radar, quiero saber si una publicación es nueva, vigente, modificada o retirada "
        "para no analizar inventario desactualizado.",
        "Evita invertir tiempo en publicaciones que ya no están disponibles o cambiaron de condiciones.",
        "Cada publicación muestra un estado temporal coherente y conserva un historial útil.",
        [
            "Aplicar la configuración de ciudad, precio, área, estrato y fecha de revisión.",
            "Registrar eventos como nuevo, visto, precio actualizado, retirado, expirado o descartado.",
            "Aplicar reglas de vigencia y una cadencia de revisión definida.",
        ],
        "Consultar una publicación y revisar su última verificación y eventos de ciclo de vida.",
        "Definir cuánto tiempo debe conservarse visible el histórico de retirados y expirados.",
        phase=1,
    )

    doc.add_page_break()
    add_hu_catalog_block(
        doc,
        bullet_id,
        "HU-F1-03",
        "Motor de activos bancarios",
        "Como inversionista, quiero comparar activos bancarios contra el mercado abierto aunque el banco "
        "no publique todos los atributos, para detectar oportunidades sin descartar inventario útil.",
        "Amplía el universo de oportunidades y trata de forma responsable la información incompleta.",
        "El activo se normaliza, muestra su frescura y se compara cuando existe evidencia suficiente.",
        [
            "Normalizar datos provenientes de diferentes bancos.",
            "Tratar atributos ausentes, como el estrato, como datos opcionales.",
            "Evitar comparaciones inválidas cuando faltan área o ubicación confiable.",
            "Conservar identidad, cambio de precio, retiro e historial del activo.",
        ],
        "Abrir una ficha bancaria y revisar la información disponible, la frescura y su comparación.",
        "Aprobar los atributos mínimos que permiten publicar una ficha bancaria incompleta.",
        phase=1,
    )
    doc.add_page_break()
    add_hu_catalog_block(
        doc,
        bullet_id,
        "HU-F1-04",
        "Motor de activos Grupo AVAL",
        "Como inversionista, quiero consultar activos AVAL extraídos de sus boletines PDF y conocer sus "
        "cambios para analizarlos junto con las demás oportunidades.",
        "Convierte una fuente documental difícil de consultar en inventario estructurado y comparable.",
        "Cada registro del boletín se transforma en una ficha deduplicada, normalizada y trazable.",
        [
            "Resolver, descargar y leer el boletín PDF vigente.",
            "Extraer código, precio, tipo, área, dirección, ciudad y datos disponibles.",
            "Deduplicar registros y comparar boletines para identificar novedades y cambios.",
            "Marcar datos ambiguos para revisión en vez de completar información por suposición.",
        ],
        "Seleccionar un activo AVAL y rastrear el registro del boletín del que proviene.",
        "Definir si las marcas “Nuevo inmueble” y “Nuevo valor” del boletín son autoritativas o deben recalcularse.",
        phase=1,
    )

    doc.add_page_break()
    add_hu_catalog_block(
        doc,
        bullet_id,
        "HU-F1-05",
        "Motor de remates judiciales",
        "Como inversionista, quiero distinguir el riesgo jurídico, la postura mínima y el porcentaje real "
        "del bien rematado para evitar interpretar una subasta de forma equivocada.",
        "Reduce el riesgo de confundir el porcentaje de postura con la participación real sobre el inmueble.",
        "La ficha resume tipo de proceso, demandante, cuota-parte, postura y condiciones de acceso.",
        [
            "Identificar el origen del demandante y el tipo de proceso.",
            "Distinguir porcentaje de postura y porcentaje del dominio rematado.",
            "Aplicar una regla conservadora cuando la fuente sea ambigua.",
            "Advertir anomalías y evitar presentar evidencia visual engañosa.",
        ],
        "Abrir un remate y revisar su lectura jurídica, cuota-parte, postura y alertas.",
        "Aprobar qué evidencia visual o documental permite publicar una ficha sin imagen original.",
        phase=1,
    )
    add_phase_close(
        doc,
        "Resultado de la Fase 1",
        "El usuario recibe un inventario más confiable y una lectura comparable de oportunidades provenientes "
        "de fuentes con estructuras, niveles de detalle y riesgos diferentes.",
    )

    doc.add_page_break()
    ui.add_heading(doc, "4. Catálogo funcional — Fase 2", 1)
    ui.add_body(
        doc,
        "Las historias de Fase 2 convierten el motor analítico en una experiencia de producto: acceso, "
        "segmentación por planes, seguimiento, monetización, operación y nuevas capas de análisis.",
    )

    add_hu_catalog_block(
        doc,
        bullet_id,
        "HU-F2-01",
        "Inicio de sesión con Google",
        "Como usuario, quiero ingresar con mi cuenta de Google para acceder al Radar sin crear otra contraseña.",
        "Disminuye la fricción de registro y facilita el retorno al producto.",
        "El usuario autoriza su cuenta, vuelve al Radar y conserva una sesión segura.",
        [
            "Mostrar el acceso con Google en la pantalla de ingreso.",
            "Completar OAuth y crear o reconocer al usuario correspondiente.",
            "Mantener la sesión y aplicar las políticas de acceso definidas.",
        ],
        "Completar el ingreso con Google y regresar al producto con la cuenta identificada.",
        "Confirmar si cualquier cuenta externa puede registrarse o si existirán dominios o invitaciones permitidas.",
        phase=2,
    )
    doc.add_page_break()
    add_hu_catalog_block(
        doc,
        bullet_id,
        "HU-F2-02",
        "Cuenta, planes y acceso freemium/Pro",
        "Como usuario, quiero conocer mi plan y acceder a beneficios acordes con mi suscripción para "
        "entender el valor de pasar a Pro.",
        "Hace visible la propuesta comercial y evita que las restricciones parezcan errores del producto.",
        "El contenido, los límites y las llamadas a acción cambian de manera coherente según el plan.",
        [
            "Mostrar cuenta, plan actual y beneficios disponibles.",
            "Restringir únicamente la información definida como premium.",
            "Aplicar los permisos tanto en la interfaz como en servidor y base de datos.",
            "Registrar interés, activación y cancelación de manera trazable.",
        ],
        "Comparar la experiencia de un visitante, un usuario registrado y un usuario Pro.",
        "Aprobar la matriz exacta de funcionalidades, límites y condiciones de cada plan.",
        phase=2,
    )

    doc.add_page_break()
    add_hu_catalog_block(
        doc,
        bullet_id,
        "HU-F2-03",
        "Alertas por correo",
        "Como usuario, quiero recibir alertas por correo cuando aparezcan oportunidades compatibles para "
        "reaccionar a tiempo.",
        "Convierte una búsqueda manual recurrente en un seguimiento automático de oportunidades.",
        "El usuario define preferencias y recibe un correo claro, relevante y sin duplicados.",
        [
            "Asociar la alerta a ciudades, zonas, tipos y criterios del usuario.",
            "Enviar un correo profesional con enlace a la oportunidad.",
            "Evitar duplicados y registrar entrega, error o rebote.",
            "Permitir ajustar o desactivar la frecuencia de notificación.",
        ],
        "Crear una alerta de prueba y revisar el correo recibido y su enlace al Radar.",
        "Definir frecuencia, horarios, umbrales y reglas de agrupación de alertas.",
        phase=2,
    )
    doc.add_page_break()
    add_hu_catalog_block(
        doc,
        bullet_id,
        "HU-F2-04",
        "Administración y operación",
        "Como administrador, quiero revisar usuarios, intereses y estado operativo para controlar el servicio.",
        "Permite operar el producto sin depender de consultas manuales a la base de datos.",
        "Un administrador autorizado puede revisar información comercial y técnica con trazabilidad.",
        [
            "Restringir el acceso administrativo por rol.",
            "Consultar usuarios, planes e intereses comerciales.",
            "Revisar indicadores de fuentes, scraping, alertas y errores.",
            "Registrar las acciones administrativas relevantes.",
        ],
        "Ingresar con un rol administrativo y revisar usuarios, intereses y estado de operación.",
        "Aprobar los roles internos y qué acciones requieren auditoría o doble validación.",
        phase=2,
    )

    doc.add_page_break()
    add_hu_catalog_block(
        doc,
        bullet_id,
        "HU-F2-05",
        "Comparables de arriendo y rentabilidad",
        "Como inversionista, quiero estimar el canon de arriendo de una propiedad y su rentabilidad con "
        "comparables similares para tomar una decisión más completa.",
        "Complementa la lectura de compra con una expectativa realista de ingreso y retorno.",
        "El análisis muestra muestra comparable, canon estimado, rentabilidad y nivel de confianza.",
        [
            "Separar correctamente ofertas de venta y arriendo.",
            "Comparar inmuebles por zona, tipo, área y atributos relevantes.",
            "Exigir una muestra mínima y mostrar la cobertura disponible.",
            "Evitar publicar una estimación cuando la confianza sea insuficiente.",
        ],
        "Analizar una propiedad de venta y revisar sus comparables de arriendo y rentabilidad estimada.",
        "Aprobar ciudades, tipos de inmueble, muestra mínima y forma de comunicar la confianza.",
        phase=2,
    )
    doc.add_page_break()
    add_hu_catalog_block(
        doc,
        bullet_id,
        "HU-F2-06",
        "Pago y activación de plan",
        "Como usuario, quiero pagar un plan de forma segura para activar los beneficios Pro.",
        "Conecta la propuesta de valor con una activación comercial verificable.",
        "Una transacción aprobada activa el plan una sola vez y queda conciliada con el proveedor de pagos.",
        [
            "Crear una referencia única de pago y mostrar el checkout.",
            "Validar el resultado mediante un webhook firmado.",
            "Activar el plan de forma idempotente y registrar la conciliación.",
            "Manejar pagos rechazados, pendientes, cancelaciones y reembolsos.",
        ],
        "Completar una transacción Sandbox y verificar la activación de la cuenta.",
        "La integración con Wompi debe retomarse únicamente con credenciales y datos legítimos del comercio.",
        phase=2,
    )

    doc.add_page_break()
    add_hu_catalog_block(
        doc,
        bullet_id,
        "HU-F2-07",
        "Filtros y navegación responsive",
        "Como usuario, quiero filtrar oportunidades con controles claros en desktop y móvil para encontrar "
        "propiedades con menos esfuerzo.",
        "Reduce carga cognitiva y facilita comparar resultados sin perder el contexto de búsqueda.",
        "Los filtros son visibles en desktop, accesibles en móvil y reflejan claramente el estado aplicado.",
        [
            "Aprovechar el panel lateral disponible en pantallas de escritorio.",
            "Adaptar filtros y resultados a móvil sin ocultar acciones principales.",
            "Mostrar filtros activos y permitir limpiarlos fácilmente.",
            "Mantener orden, paginación y contexto al abrir o cerrar una ficha.",
        ],
        "Aplicar filtros en desktop y móvil, abrir una propiedad y regresar conservando la búsqueda.",
        "Validar comprensión, navegación por teclado y criterios WCAG con usuarios representativos.",
        phase=2,
    )
    add_phase_close(
        doc,
        "Resultado de la Fase 2",
        "El motor de oportunidades se convierte en un servicio utilizable y operable: el usuario puede "
        "acceder, filtrar, seguir, analizar y contratar; el equipo puede administrar y observar el producto.",
    )

    doc.add_page_break()
    ui.add_heading(doc, "5. Criterios transversales de aceptación", 1)
    ui.add_body(
        doc,
        "Además de los criterios particulares de cada HU, el producto necesita condiciones comunes que "
        "protejan la confianza del usuario y la operación en producción.",
    )
    for text, prefix in [
        ("Calidad de datos: no completar información incierta por suposición; marcar y aislar anomalías.", "Calidad de datos:"),
        ("Trazabilidad: conservar fuente, fecha de verificación y eventos relevantes.", "Trazabilidad:"),
        ("Seguridad: aplicar autenticación, roles y permisos también en servidor y base de datos.", "Seguridad:"),
        ("Privacidad: usar datos personales solo para cuenta, alertas y operación autorizada.", "Privacidad:"),
        ("Experiencia: explicar restricciones, estados vacíos, errores y acciones siguientes.", "Experiencia:"),
        ("Accesibilidad: permitir lectura, contraste, teclado y adaptación responsive.", "Accesibilidad:"),
        ("Operación: registrar errores, reintentos, duplicados y resultados de procesos automáticos.", "Operación:"),
    ]:
        ui.add_list_item(doc, text, bullet_id, bold_prefix=prefix)

    ui.add_heading(doc, "6. Validaciones que debe resolver el cliente", 1)
    for text in [
        "Aprobar la redacción, actor y valor de cada historia de usuario.",
        "Confirmar la regla geográfica del motor de comparables.",
        "Definir retención de publicaciones retiradas o expiradas.",
        "Aprobar atributos mínimos para activos bancarios incompletos.",
        "Definir el tratamiento de novedades en boletines AVAL.",
        "Aprobar la evidencia mínima para remates sin imagen original.",
        "Definir funcionalidades y límites de los planes gratuito y Pro.",
        "Acordar frecuencia, preferencias y agrupación de alertas.",
        "Priorizar ciudades y muestra mínima para comparables de arriendo.",
        "Confirmar el aplazamiento de Wompi hasta contar con información real del comercio.",
    ]:
        ui.add_list_item(doc, text, bullet_id)

    doc.add_page_break()
    ui.add_heading(doc, "7. Forma recomendada de trabajar las HU", 1)
    for text in [
        "Validar y aprobar el catálogo funcional con el cliente.",
        "Asignar prioridad y responsable a cada HU.",
        "Convertir los criterios principales en casos de aceptación demostrables.",
        "Registrar cualquier solicitud nueva con un ID y su impacto en alcance.",
        "Cerrar cada HU con demostración, evidencia y aprobación, no solo con código implementado.",
    ]:
        ui.add_list_item(doc, text, workflow_decimal_id)
    ui.add_callout(
        doc,
        "Regla de control",
        "Ninguna funcionalidad nueva debería incorporarse al alcance sin actor, historia, valor, criterios "
        "de aceptación, prioridad y validación del cliente.",
        fill=ui.GOLD_LIGHT,
        accent=ui.AMBER,
    )

    doc.add_page_break()
    ui.add_heading(doc, "8. Documentos y decisiones de referencia", 1)
    for text in [
        "Radar de Oportunidades — Documentación y Propuesta.",
        "Radar CRECE — Índice Maestro HU Fase 1.",
        "Radar CRECE — Especificación Motor de Comparables v1.2.",
        "HU Núcleo Motor de Comparables.",
        "HU Rango de Precio y Vigencia del Motor Radar.",
        "HU Motor de Activos Bancarios.",
        "HU Motor de Activos Bancarios — Grupo AVAL.",
        "HU Motor de Remates Judiciales.",
        "Hoja de ruta del documento inicial: Google, freemium, alertas, panel, comparador y rentabilidad.",
        "Decisiones posteriores: Resend, Wompi, filtros laterales y comparables de arriendo.",
    ]:
        ui.add_list_item(doc, text, bullet_id)

    ui.add_callout(
        doc,
        "Uso recomendado",
        "Utilizar este documento como base de conversación y convertir las decisiones aprobadas en una "
        "matriz de aceptación firmada o en un backlog compartido con control de cambios.",
        fill=ui.PURPLE_LIGHT,
        accent=ui.PURPLE,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
