from __future__ import annotations

from datetime import date
from pathlib import Path
import subprocess

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "output" / "docs" / "Matriz_Trazabilidad_Historias_Usuario_Radar_2026-07-25.docx"

# standard_business_brief preset + named Radar CRECE brand override.
FONT = "Calibri"
INK = "2A1336"
PURPLE = "4F2868"
PURPLE_LIGHT = "F4EFF7"
GOLD = "F2CA04"
GOLD_LIGHT = "FFF9D9"
BLUE = "2E74B5"
MUTED = "655B69"
GRAY = "F2F4F7"
MID_GRAY = "D9DDE3"
GREEN = "1E7A46"
GREEN_LIGHT = "EAF6EF"
AMBER = "8A6300"
AMBER_LIGHT = "FFF5D6"
RED = "9B1C1C"
RED_LIGHT = "FDECEC"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_font(run, size=11, color=INK, bold=False, italic=False, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, spec in edges.items():
        tag = f"w:{edge_name}"
        edge = tc_borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            tc_borders.append(edge)
        for key in ("val", "sz", "color", "space"):
            if key in spec:
                edge.set(qn(f"w:{key}"), str(spec[key]))


def set_table_geometry(table, widths, indent=TABLE_INDENT):
    assert sum(widths) == TABLE_WIDTH, (widths, sum(widths))
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_font(run, size=9, color=MUTED)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, text, fld_end])


def add_numbering(doc, kind):
    numbering = doc.part.numbering_part.element
    existing = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(existing, default=0) + 1
    nums = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(nums, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{abstract_id:08X}")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(r_fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(num)
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), bold=True)
        set_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_font(p.add_run(text))
    return p


def add_body(doc, text, bold_prefix=None, italic=False, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), bold=True, color=color)
        set_font(p.add_run(text[len(bold_prefix):]), italic=italic, color=color)
    else:
        set_font(p.add_run(text), italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    set_font(r, size=9, color=AMBER, bold=True)
    r.font.all_caps = True
    return p


def add_callout(doc, label, text, fill=PURPLE_LIGHT, accent=PURPLE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    set_cell_border(
        cell,
        start={"val": "single", "sz": "22", "color": accent, "space": "0"},
        top={"val": "nil"},
        bottom={"val": "nil"},
        end={"val": "nil"},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(f"{label}: "), size=11, color=accent, bold=True)
    set_font(p.add_run(text), size=11, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_status_line(doc, progress, state, tone="green"):
    colors = {
        "green": (GREEN_LIGHT, GREEN),
        "amber": (AMBER_LIGHT, AMBER),
        "red": (RED_LIGHT, RED),
    }
    fill, color = colors[tone]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [1500, 7860])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, fill)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": fill},
            bottom={"val": "single", "sz": "4", "color": fill},
            start={"val": "single", "sz": "4", "color": fill},
            end={"val": "single", "sz": "4", "color": fill},
        )
    p0 = table.cell(0, 0).paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p0.add_run(progress), size=12, color=color, bold=True)
    p1 = table.cell(0, 1).paragraphs[0]
    set_font(p1.add_run(state), size=10.5, color=color, bold=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def add_summary_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [1450, 4050, 1000, 2860]
    set_table_geometry(table, widths)
    headers = ["ID", "Historia / bloque funcional", "Avance", "Lectura ejecutiva"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, INK)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
        set_font(p.add_run(text), size=9.5, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])

    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(text), size=9.2, color=INK, bold=(idx in (0, 2)))
        pct = int(row_data[2].replace("%", "").replace(",", ".").split(".")[0])
        set_cell_shading(cells[2], GREEN_LIGHT if pct >= 85 else AMBER_LIGHT if pct >= 70 else RED_LIGHT)
    return table


def add_scope_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [2750, 1420, 1850, 3340]
    set_table_geometry(table, widths)
    headers = ["Funcionalidad añadida", "Estado", "Clasificación", "Paso siguiente"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, PURPLE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
        set_font(p.add_run(text), size=9.3, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    for feature, state, classification, next_step in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((feature, state, classification, next_step)):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(text), size=9, bold=(idx == 0))
        if state in {"Operativo", "Desplegado", "Validado"}:
            set_cell_shading(cells[1], GREEN_LIGHT)
        else:
            set_cell_shading(cells[1], AMBER_LIGHT)
    return table


def add_next_steps_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [750, 2550, 2900, 3160]
    set_table_geometry(table, widths)
    headers = ["Prio.", "Trabajo", "Entregable verificable", "Criterio de salida"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, INK)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
        set_font(p.add_run(text), size=9.2, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    for priority, work, deliverable, exit_criteria in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((priority, work, deliverable, exit_criteria)):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(text), size=8.9, bold=(idx in (0, 1)))
        set_cell_shading(cells[0], RED_LIGHT if priority == "P0" else AMBER_LIGHT if priority == "P1" else GRAY)
    return table


def configure_document(doc):
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        1: (16, PURPLE, 16, 8),
        2: (13, PURPLE, 12, 6),
        3: (12, INK, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(3)
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [5200, 4160], indent=0)
    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(left.add_run("RADAR CRECE  |  HISTORIAS DE USUARIO"), size=8.5, color=MUTED, bold=True)
    set_font(right.add_run("Documento de avance"), size=8.5, color=MUTED)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=40, start=0, end=0)
        set_cell_border(cell, bottom={"val": "single", "sz": "5", "color": MID_GRAY, "space": "0"})

    footer = section.footer
    fp = footer.paragraphs[0]
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(footer_table, [6500, 2860], indent=0)
    l = footer_table.cell(0, 0).paragraphs[0]
    r = footer_table.cell(0, 1).paragraphs[0]
    set_font(l.add_run("Corte: 25 de julio de 2026  •  Documento de trabajo para validación"), size=8.5, color=MUTED)
    add_page_field(r)
    for cell in footer_table.rows[0].cells:
        set_cell_margins(cell, top=40, bottom=0, start=0, end=0)
        set_cell_border(cell, top={"val": "single", "sz": "5", "color": MID_GRAY, "space": "0"})

    core = doc.core_properties
    core.title = "Matriz de trazabilidad de historias de usuario — Radar de Oportunidades"
    core.subject = "Avance funcional, evidencia, pendientes y próximos pasos"
    core.author = "Equipo de desarrollo Radar CRECE"
    core.keywords = "Radar, CRECE, historias de usuario, trazabilidad, avance, fase 1, fase 2"
    core.comments = "Documento preparado para revisión y validación con el cliente."


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340])
    metrics = [
        ("84%", "Fase 1 defendible"),
        ("141 / 141", "Pruebas aprobadas"),
        ("10 / 10", "Recorridos E2E locales"),
        ("0", "Errores TypeScript"),
    ]
    for idx, (value, label) in enumerate(metrics):
        cell = table.cell(0, idx)
        set_cell_shading(cell, PURPLE_LIGHT if idx != 0 else GOLD_LIGHT)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "12", "color": GOLD if idx == 0 else PURPLE, "space": "0"},
            bottom={"val": "single", "sz": "4", "color": MID_GRAY, "space": "0"},
            start={"val": "single", "sz": "4", "color": MID_GRAY, "space": "0"},
            end={"val": "single", "sz": "4", "color": MID_GRAY, "space": "0"},
        )
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(value), size=16, color=INK, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        set_font(p2.add_run(label), size=8.5, color=MUTED, bold=True)


def add_hu_block(doc, bullet_id, hu_id, title, progress, state, tone, objective, done, evidence, pending, decision):
    add_heading(doc, f"{hu_id} — {title}", 2)
    add_status_line(doc, progress, state, tone)
    add_body(doc, f"Objetivo. {objective}", bold_prefix="Objetivo. ")
    add_body(doc, "Implementado:", bold_prefix="Implementado:")
    for item in done:
        add_list_item(doc, item, bullet_id)
    add_body(doc, f"Evidencia verificable. {evidence}", bold_prefix="Evidencia verificable. ")
    add_body(doc, f"Pendiente para cierre. {pending}", bold_prefix="Pendiente para cierre. ")
    add_callout(doc, "Decisión o validación del cliente", decision, fill=GOLD_LIGHT, accent=AMBER)


def main():
    doc = Document()
    configure_document(doc)
    bullet_id = add_numbering(doc, "bullet")
    decimal_id = add_numbering(doc, "decimal")

    try:
        commit = subprocess.check_output(
            ["git", "rev-parse", "--short", "HEAD"], cwd=ROOT, text=True
        ).strip()
    except Exception:
        commit = "no disponible"

    # Customer-pack opening.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run("INFORME DE AVANCE Y TRAZABILIDAD"), size=9.5, color=AMBER, bold=True)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    set_font(title.add_run("Radar de Oportunidades Inmobiliarias"), size=28, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_font(
        subtitle.add_run("Matriz de historias de usuario, evidencia y próximos pasos"),
        size=14,
        color=PURPLE,
        bold=True,
    )

    meta = doc.add_table(rows=2, cols=2)
    set_table_geometry(meta, [4680, 4680], indent=0)
    metadata = [
        ("Preparado para", "Validación con el cliente"),
        ("Corte funcional", "25 de julio de 2026"),
        ("Versión de código", commit),
        ("Entorno", "Local validado y producción activa"),
    ]
    for idx, (label, value) in enumerate(metadata):
        cell = meta.cell(idx // 2, idx % 2)
        set_cell_shading(cell, GRAY)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": MID_GRAY},
            bottom={"val": "single", "sz": "4", "color": MID_GRAY},
            start={"val": "single", "sz": "4", "color": MID_GRAY},
            end={"val": "single", "sz": "4", "color": MID_GRAY},
        )
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(label.upper()), size=8, color=MUTED, bold=True)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        set_font(p2.add_run(value), size=10.5, color=INK, bold=True)

    doc.add_paragraph()
    add_metric_strip(doc)
    doc.add_paragraph()
    add_callout(
        doc,
        "Conclusión ejecutiva",
        "Las historias de usuario originales sí han sido la base del motor y de las reglas de negocio. "
        "La Fase 1 alcanza un 83,6%, comunicado como 84%. El producto es un MVP funcional y demostrable; "
        "los pendientes están identificados y se concentran en trazabilidad completa del ciclo de vida, "
        "normalización de fuentes, operación de producción y decisiones de alcance.",
        fill=PURPLE_LIGHT,
        accent=PURPLE,
    )
    add_body(
        doc,
        "Este documento separa deliberadamente las historias originales de las funcionalidades solicitadas "
        "después. Así se evita presentar ampliaciones como requisitos cerrados o sumar trabajo adicional al "
        "porcentaje base sin aprobación del cliente.",
        italic=True,
        color=MUTED,
    )
    add_body(
        doc,
        "URL de referencia: https://joinsclee-radar.juno8i.easypanel.host/",
        bold_prefix="URL de referencia: ",
    )

    doc.add_page_break()
    add_heading(doc, "1. Cómo leer esta matriz", 1)
    add_body(
        doc,
        "La evaluación se realiza por bloque funcional documentado. El porcentaje representa cumplimiento "
        "del alcance de la Fase 1 y no significa que cada fuente o proceso esté al 100% de madurez operativa."
    )
    for text in [
        "Implementado: existe comportamiento funcional verificable en código o interfaz.",
        "Evidencia: existen pruebas, rutas, reglas o datos que permiten demostrar el resultado.",
        "Pendiente: falta una condición de aceptación, cobertura de datos o control operativo.",
        "Decisión del cliente: definición de negocio necesaria para cerrar la historia sin asumir reglas.",
        "Alcance adicional: funcionalidad solicitada después de la línea base y que debe versionarse como nueva HU o extensión.",
    ]:
        add_list_item(doc, text, bullet_id)

    add_heading(doc, "2. Resumen de cumplimiento", 1)
    add_summary_table(
        doc,
        [
            ("HU-F1-01", "Núcleo de comparables e Índice CRECE", "95%", "Prácticamente cerrada; requiere auditoría histórica y definición geográfica."),
            ("HU-F1-02", "Rango de precio y vigencia del Portal", "78%", "Reglas implementadas; falta uniformar eventos del ciclo de vida."),
            ("HU-F1-03", "Motor de activos bancarios", "82%", "Operativo; falta identidad secundaria e historial uniforme."),
            ("HU-F1-04", "Motor de activos Grupo AVAL", "72%", "Ingesta PDF funcional; normalización y comparación entre boletines pendientes."),
            ("HU-F1-05", "Motor de remates judiciales", "92%", "Reglas jurídicas y comerciales sólidas; falta cuarentena de anomalías."),
            ("PRD-01", "UX, autenticación y modelo freemium", "88%", "Experiencia funcional y responsive; quedan controles de activación."),
            ("OPS-01", "Producción, seguridad y operación", "64%", "Producción activa; endurecimiento y recuperación aún en curso."),
            ("TOTAL", "Fase 1 ponderada", "83,6%", "Se comunica como 84%; supera el objetivo de demostración del 80%."),
        ],
    )
    add_body(
        doc,
        "Nota metodológica. Los porcentajes son una ponderación por bloques funcionales. No se deben promediar "
        "de forma simple ni usar como certificado de producción masiva.",
        bold_prefix="Nota metodológica. ",
        color=MUTED,
    )

    doc.add_page_break()
    add_heading(doc, "3. Trazabilidad detallada de las historias originales", 1)
    add_hu_block(
        doc,
        bullet_id,
        "HU-F1-01",
        "Núcleo de comparables e Índice CRECE",
        "95%",
        "Avanzada — en cierre funcional",
        "green",
        "Comparar cada inmueble contra una muestra de mercado abierto equivalente y traducir el resultado a una categoría CRECE comprensible.",
        [
            "Mediana robusta por precio por metro cuadrado y recorte de valores atípicos.",
            "Comparación por tipo, área, atributos y cascada geográfica; muestra mínima controlada.",
            "Exclusión de bancos y remates del universo de referencia del mercado abierto.",
            "Tabla maestra de once categorías CRECE y tratamiento definido del rango 0,71–0,75.",
            "Persistencia de índice/categoría y control de acceso desde el servidor.",
        ],
        "Motor TypeScript de comparables y clasificación, criterios trazables mostrados al usuario y pruebas de límites de categoría y casos de referencia.",
        "Auditar la persistencia del Índice CRECE en el 100% del histórico y documentar el criterio definitivo de radio o zonas colindantes.",
        "Aprobar cuándo el motor puede ampliar de barrio a zona/ciudad y cuál debe ser el radio máximo aceptable.",
    )
    doc.add_page_break()
    add_hu_block(
        doc,
        bullet_id,
        "HU-F1-02",
        "Rango de precio y vigencia del Portal",
        "78%",
        "Parcial — reglas listas, trazabilidad incompleta",
        "amber",
        "Controlar qué ciudades y rangos se monitorean, identificar cambios y evitar que publicaciones antiguas parezcan inventario vigente.",
        [
            "Zonas monitoreadas configurables por ciudad y filtros de precio, área, estrato y fecha.",
            "Reglas de 30 y 90 días implementadas y probadas.",
            "Modelo de eventos de vigencia y cadencia semanal disponible.",
            "Filtros de interfaz coherentes con los atributos capturados.",
        ],
        "Configuración persistente de zonas, motor de vigencia y pruebas sobre estados temporales.",
        "Conectar uniformemente los eventos NUEVO, VISTO, PRECIO_ACTUALIZADO, RETIRADO, EXPIRADO y DESCARTADO desde todos los scrapers; verificar activación y población de ciudades.",
        "Definir si un inmueble retirado debe conservarse visible como histórico y durante cuánto tiempo.",
    )

    doc.add_page_break()
    add_hu_block(
        doc,
        bullet_id,
        "HU-F1-03",
        "Motor de activos bancarios",
        "82%",
        "Avanzada — cierre de identidad e historial",
        "amber",
        "Normalizar activos de diferentes bancos y determinar si son oportunidades contra el mercado abierto sin depender de campos que algunas fuentes no publican.",
        [
            "Fuentes bancarias normalizadas y clasificación contra comparables del mercado abierto.",
            "Ausencia de estrato tratada como dato opcional, no como bloqueo.",
            "Mensaje de última verificación y rotación semanal determinista del inventario.",
            "Interfaz y filtros adaptados a la calidad real de los datos bancarios.",
        ],
        "Scrapers bancarios, motor de comparables, rotación semanal y pruebas de frescura.",
        "Cerrar la identidad secundaria para enlaces o slugs inestables; uniformar detección de retiro, cambio de precio e historial entre bancos.",
        "Confirmar qué atributos mínimos debe tener un activo bancario para publicarse aun cuando la fuente no entregue estrato, habitaciones o baños.",
    )
    doc.add_page_break()
    add_hu_block(
        doc,
        bullet_id,
        "HU-F1-04",
        "Motor de activos Grupo AVAL",
        "72%",
        "Parcial — ingesta funcional, normalización pendiente",
        "amber",
        "Extraer y comparar activos publicados en boletines PDF de AVAL conservando trazabilidad del documento fuente.",
        [
            "Resolución dinámica, descarga y lectura por página del boletín PDF.",
            "Deduplicación por código y extracción de precio, tipo, área, dirección y ciudad.",
            "Campo area_tipo para evitar comparaciones inválidas entre lotes y vivienda.",
            "Imagen por ficha y soporte de almacenamiento.",
        ],
        "Parser PDF, deduplicación, modelo normalizado y representación en la aplicación.",
        "Normalización ciudad–departamento con bandera de revisión; campos opcionales de habitaciones/baños; comparación entre boletines e insignias de novedad/cambio; trazabilidad completa del boletín.",
        "Aprobar si las insignias “Nuevo inmueble” y “Nuevo valor” son fuente suficiente o deben recalcularse comparando boletines.",
    )

    doc.add_page_break()
    add_hu_block(
        doc,
        bullet_id,
        "HU-F1-05",
        "Motor de remates judiciales",
        "92%",
        "Avanzada — cierre de calidad de datos",
        "green",
        "Clasificar el riesgo y el acceso comercial de cada remate sin confundir postura mínima, avalúo y porcentaje real del bien.",
        [
            "Identificación del origen del demandante, tipo de proceso y cuota-parte.",
            "Regla conservadora cuando el origen es incierto.",
            "Matriz gratis/suscripción y alerta visible para riesgos jurídicos.",
            "Diferenciación entre porcentaje de postura y porcentaje de dominio rematado.",
            "Barreras contra valores financieros y porcentajes imposibles antes de mostrarlos.",
        ],
        "Motor jurídico, matriz de acceso, sanitización de datos y pruebas específicas sobre avisos y porcentajes.",
        "Enviar anomalías a cuarentena desde ingestión/base de datos y reemplazar imágenes genéricas por evidencia neutral y trazable cuando la fuente no entregue imagen.",
        "Definir qué evidencia visual o documental es suficiente para publicar una ficha sin imagen original.",
    )

    doc.add_page_break()
    add_heading(doc, "4. Requisitos transversales de producto y operación", 1)
    add_body(
        doc,
        "Estos bloques no sustituyen las HU del motor, pero determinan si el usuario puede utilizar el producto "
        "de forma segura, comprensible y repetible."
    )
    add_heading(doc, "PRD-01 — UX, autenticación y freemium", 2)
    add_status_line(doc, "88%", "Avanzado — experiencia demostrable", "green")
    for item in [
        "Navegación responsive y accesible con filtros laterales en escritorio y panel desplegable en móvil.",
        "Tarjetas, modales, estados de carga, foco, teclado y targets táctiles mejorados.",
        "Favoritos anónimos y sincronización al iniciar sesión.",
        "Registro por correo y acceso con Google OAuth disponible; el proveedor externo se validó manualmente.",
        "Valor visible antes del registro: comparables, costos y señales CRECE.",
    ]:
        add_list_item(doc, item, bullet_id)
    add_body(
        doc,
        "Pendiente para cierre. Onboarding posterior al registro, activación controlada de alertas recurrentes y revisión formal de accesibilidad.",
        bold_prefix="Pendiente para cierre. ",
    )
    add_heading(doc, "OPS-01 — Producción, seguridad y operación", 2)
    add_status_line(doc, "64%", "En endurecimiento — producción activa, operación no cerrada", "amber")
    for item in [
        "HTTPS, healthcheck, readiness, apagado controlado, cabeceras de seguridad y request IDs.",
        "Imagen Docker multi-stage y ejecución no-root.",
        "CI, pruebas automatizadas y monitoreo sintético.",
        "Runbook, backup lógico y verificación de artefactos.",
    ]:
        add_list_item(doc, item, bullet_id)
    add_body(
        doc,
        "Pendiente prioritario. Sesiones en cookies HttpOnly/Secure/SameSite, verificación de correo, reducción de datos personales en logs, centralización de métricas/trazas, restauración real en entorno de ensayo y prueba de rollback.",
        bold_prefix="Pendiente prioritario. ",
    )

    doc.add_page_break()
    add_heading(doc, "5. Alcance adicional solicitado después de la línea base", 1)
    add_callout(
        doc,
        "Regla de alcance",
        "Estos elementos agregan valor real, pero deben registrarse como nuevas HU o extensiones. "
        "No se usan para inflar el 84% de Fase 1.",
        fill=GOLD_LIGHT,
        accent=AMBER,
    )
    add_scope_table(
        doc,
        [
            ("Rediseño visual y filtros laterales", "Desplegado", "Extensión UX", "Validación final con usuarios y revisión WCAG."),
            ("Inicio de sesión con Google", "Validado manualmente", "Extensión Auth", "Monitorear altas externas y políticas de acceso."),
            # "Operativo" pinta la celda en verde y afirma que salen correos. El canal
            # está construido y probado, pero el trabajo `alertas` de `radar_cron_jobs`
            # sigue deshabilitado por decisión de producto: hoy no se envía ninguno.
            ("Alertas por correo con Resend", "Configurado, envío en pausa", "Nueva HU", "Ejecutar el canary dirigido y habilitar el despachador semanal cuando el responsable del producto lo apruebe."),
            ("Planes, cuenta, comparador y panel admin", "Demo funcional", "Fase 2 comercial", "Validar prioridades, permisos y operación con usuarios piloto."),
            ("Wompi Web Checkout", "Preparado", "Fase 2 / tercero", "Retomar cuando existan credenciales legítimas y ejecutar compra Sandbox E2E."),
            ("Comparables de arriendo y rentabilidad", "Piloto", "Nueva HU", "Consolidar cobertura, calidad y confianza del canon observado."),
        ],
    )
    add_heading(doc, "5.1 Piloto de arriendos", 2)
    add_body(
        doc,
        "El flujo ya separa el mercado de venta del mercado de arriendo, almacena cánones mensuales, calcula "
        "una referencia de alquiler para propiedades similares y permite recalcular rentabilidad. El piloto "
        "controlado cubre Medellín, Bogotá, Cali, Barranquilla y Bucaramanga."
    )
    add_body(
        doc,
        "Pendiente. Ejecutar ingestas periódicas completas, medir cobertura por ciudad/tipo, definir una muestra mínima y mostrar un nivel de confianza que impida presentar un canon débil como dato concluyente.",
        bold_prefix="Pendiente. ",
    )

    doc.add_page_break()
    add_heading(doc, "6. Próximos pasos priorizados", 1)
    add_next_steps_table(
        doc,
        [
            ("P0", "Validar esta matriz con el cliente", "Versión aprobada de HU, criterios y alcance adicional.", "Cada historia tiene aceptación, prioridad y responsable definidos."),
            ("P0", "Cerrar ciclo de vida Portal/Bancos", "Eventos uniformes y auditables desde todos los scrapers.", "Muestra de publicaciones con historial completo y estados reproducibles."),
            ("P1", "Cerrar normalización AVAL", "Comparación entre boletines y reporte de datos ambiguos.", "Duplicados, cambios de precio y ciudades se resuelven de forma trazable."),
            ("P1", "Consolidar piloto de arriendos", "Reporte de cobertura, calidad, muestra y confianza por ciudad.", "El canon se muestra solo cuando supera criterios de datos acordados."),
            ("P1", "Operar alertas por correo", "Canary, agenda recurrente, métricas e idempotencia verificadas.", "Una alerta real llega una sola vez y queda registrada."),
            ("P2", "Endurecer producción", "Sesión segura, verificación de correo, logs sin PII y restauración ensayada.", "Controles críticos verificados y recuperación cronometrada."),
            ("P2", "Retomar pagos", "Transacción Wompi Sandbox aprobada, webhook e idempotencia.", "El pago aprobado activa Pro; retorno del navegador por sí solo no lo hace."),
        ],
    )

    add_heading(doc, "7. Decisiones que deben solicitarse al cliente", 1)
    for text in [
        "Radio máximo y regla para ampliar comparables desde barrio hacia zona o ciudad.",
        "Atributos mínimos para publicar activos bancarios o AVAL cuando la fuente omite campos.",
        "Ciudades, tipos de inmueble y muestra mínima para aceptar un canon de arriendo.",
        "Tratamiento visible y duración del histórico de inmuebles retirados o expirados.",
        "Separación formal entre Fase 2 técnica —remates— y Fase 2 comercial —planes, alertas, pagos y expansión—.",
        "Confirmación de que Wompi queda aplazado hasta contar con datos reales del comercio.",
    ]:
        add_list_item(doc, text, bullet_id)

    doc.add_page_break()
    add_heading(doc, "8. Guion sugerido para presentar el avance", 1)
    steps = [
        "Abrir con el resultado: la Fase 1 está en 84% y supera el objetivo de demostración del 80%.",
        "Explicar que el porcentaje sale de las historias originales, no de sumar funcionalidades nuevas.",
        "Demostrar comparables e Índice CRECE con una propiedad real y mostrar los criterios usados.",
        "Recorrer un activo bancario y un remate para evidenciar las reglas específicas de cada fuente.",
        "Mostrar filtros responsive, Google OAuth, guardados y el análisis de arriendo/rentabilidad como alcance adicional.",
        "Reconocer los pendientes sin maquillarlos: historial, AVAL, calidad de arriendos y endurecimiento operativo.",
        "Cerrar solicitando aprobación de las decisiones de la sección 7 y acordando el orden de los próximos pasos.",
    ]
    for step in steps:
        add_list_item(doc, step, decimal_id)
    add_callout(
        doc,
        "Mensaje central",
        "El producto no es una maqueta: existe una implementación demostrable y probada. "
        "Tampoco se presenta como terminado al 100%: las brechas y dependencias están identificadas.",
        fill=PURPLE_LIGHT,
        accent=PURPLE,
    )

    add_heading(doc, "9. Evidencia técnica del corte", 1)
    for item in [
        "TypeScript: compilación sin errores.",
        "Pruebas unitarias e integración: suite completa aprobada sin fallos en el corte.",
        "Recorridos E2E locales: suite crítica aprobada sin fallos en el corte.",
        "Cobertura E2E actual: dashboard/configuración, Bancos, Remates, rentabilidad con arriendos, guardados, login, disponibilidad del acceso con Google, checkout demo, móvil, filtros laterales y estados de carga.",
        f"Versión de código evaluada: {commit}.",
        "Producción: https://joinsclee-radar.juno8i.easypanel.host/",
    ]:
        add_list_item(doc, item, bullet_id)

    doc.add_page_break()
    add_heading(doc, "10. Documentos de referencia", 1)
    references = [
        "Radar de Oportunidades — Documentación y Propuesta.",
        "Radar CRECE — Índice Maestro HU Fase 1.",
        "Radar CRECE — Especificación Motor de Comparables v1.2.",
        "HU Núcleo Motor de Comparables.",
        "HU Rango de Precio y Vigencia del Motor Radar.",
        "HU Motor de Activos Bancarios.",
        "HU Motor de Activos Bancarios — Grupo AVAL.",
        "HU Motor de Remates Judiciales.",
        "Informes internos de avance de Fase 1 y Fase 2, roadmap de producción y evidencia automatizada del repositorio.",
    ]
    for ref in references:
        add_list_item(doc, ref, bullet_id)

    add_callout(
        doc,
        "Próxima actualización",
        "Después de la validación del cliente, esta matriz debe versionarse y convertirse en el backlog oficial: "
        "cada cambio nuevo deberá quedar asociado a una HU, criterio de aceptación, evidencia y fecha de aprobación.",
        fill=GOLD_LIGHT,
        accent=AMBER,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
