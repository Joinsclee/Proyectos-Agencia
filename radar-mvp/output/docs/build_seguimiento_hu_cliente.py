from __future__ import annotations

from pathlib import Path
import subprocess

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

import build_matriz_hu_cliente as ui


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "output" / "docs" / "Seguimiento_Historias_Usuario_Radar_2026-07-25.docx"


def add_hu_metric_strip(doc):
    table = doc.add_table(rows=1, cols=4)
    ui.set_table_geometry(table, [2340, 2340, 2340, 2340])
    metrics = [
        ("5", "HU trazadas en Fase 1"),
        ("84%", "Cumplimiento Fase 1"),
        ("7", "HU trazadas en Fase 2"),
        ("75%", "Avance técnico Fase 2"),
    ]
    for idx, (value, label) in enumerate(metrics):
        cell = table.cell(0, idx)
        ui.set_cell_shading(cell, ui.GOLD_LIGHT if idx == 1 else ui.PURPLE_LIGHT)
        ui.set_cell_border(
            cell,
            top={
                "val": "single",
                "sz": "12",
                "color": ui.GOLD if idx == 1 else ui.PURPLE,
                "space": "0",
            },
            bottom={"val": "single", "sz": "4", "color": ui.MID_GRAY, "space": "0"},
            start={"val": "single", "sz": "4", "color": ui.MID_GRAY, "space": "0"},
            end={"val": "single", "sz": "4", "color": ui.MID_GRAY, "space": "0"},
        )
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = ui.Pt(2)
        ui.set_font(p.add_run(value), size=16, color=ui.INK, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = ui.Pt(0)
        ui.set_font(p2.add_run(label), size=8.5, color=ui.MUTED, bold=True)


def add_acceptance_item(doc, bullet_id, status, text):
    labels = {
        "done": ("Cumplido", ui.GREEN),
        "partial": ("Parcial", ui.AMBER),
        "pending": ("Pendiente", ui.RED),
    }
    label, color = labels[status]
    p = doc.add_paragraph()
    p.paragraph_format.space_after = ui.Pt(7)
    p.paragraph_format.line_spacing = 1.167
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = ui.OxmlElement("w:ilvl")
    ilvl.set(ui.qn("w:val"), "0")
    num_pr.append(ilvl)
    num = ui.OxmlElement("w:numId")
    num.set(ui.qn("w:val"), str(bullet_id))
    num_pr.append(num)
    ui.set_font(p.add_run(f"{label}: "), color=color, bold=True)
    ui.set_font(p.add_run(text))
    return p


def add_user_story_block(
    doc,
    bullet_id,
    hu_id,
    title,
    progress,
    state,
    tone,
    story,
    criteria,
    evidence,
    pending,
    decision,
):
    ui.add_heading(doc, f"{hu_id} — {title}", 2)
    ui.add_status_line(doc, progress, state, tone)
    ui.add_callout(doc, "Historia de usuario", story, fill=ui.PURPLE_LIGHT, accent=ui.PURPLE)
    ui.add_body(doc, "Criterios de aceptación revisados:", bold_prefix="Criterios de aceptación revisados:")
    for status, criterion in criteria:
        add_acceptance_item(doc, bullet_id, status, criterion)
    ui.add_body(
        doc,
        f"Evidencia actual. {evidence}",
        bold_prefix="Evidencia actual. ",
    )
    ui.add_body(
        doc,
        f"Condición pendiente para cerrar la HU. {pending}",
        bold_prefix="Condición pendiente para cerrar la HU. ",
    )
    ui.add_callout(
        doc,
        "Validación requerida del cliente",
        decision,
        fill=ui.GOLD_LIGHT,
        accent=ui.AMBER,
    )


def add_master_hu_table(doc, rows):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    widths = [1250, 3100, 1000, 1750, 2260]
    ui.set_table_geometry(table, widths)
    headers = ["ID", "Necesidad del usuario", "Avance", "Estado", "Cierre pendiente"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        ui.set_cell_shading(cell, ui.INK)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT
        ui.set_font(p.add_run(text), size=9.1, color=ui.WHITE, bold=True)
    ui.set_repeat_table_header(table.rows[0])
    for hu_id, need, progress, state, close in rows:
        cells = table.add_row().cells
        values = (hu_id, need, progress, state, close)
        for idx, text in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = ui.Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT
            ui.set_font(p.add_run(text), size=8.8, bold=(idx in (0, 2)))
        pct = int(progress.replace("%", "").split(",")[0])
        ui.set_cell_shading(cells[2], ui.GREEN_LIGHT if pct >= 85 else ui.AMBER_LIGHT)
    return table


def add_complementary_hu_table(doc, rows):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    widths = [1250, 3210, 1000, 1500, 2400]
    ui.set_table_geometry(table, widths)
    headers = ["ID", "Necesidad del usuario", "Avance", "Estado", "Cierre pendiente"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        ui.set_cell_shading(cell, ui.PURPLE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT
        ui.set_font(p.add_run(text), size=8.9, color=ui.WHITE, bold=True)
    ui.set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = ui.Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT
            ui.set_font(p.add_run(text), size=8.5, bold=(idx == 0))
        pct = int(row[2].replace("%", ""))
        ui.set_cell_shading(cells[2], ui.GREEN_LIGHT if pct >= 85 else ui.AMBER_LIGHT)
    return table


def add_phase2_story_block(
    doc,
    bullet_id,
    hu_id,
    title,
    progress,
    state,
    tone,
    story,
    criteria,
    evidence,
    pending,
):
    ui.add_heading(doc, f"{hu_id} — {title}", 2)
    ui.add_status_line(doc, progress, state, tone)
    ui.add_callout(doc, "Historia de usuario", story, fill=ui.PURPLE_LIGHT, accent=ui.PURPLE)
    for status, criterion in criteria:
        add_acceptance_item(doc, bullet_id, status, criterion)
    ui.add_body(doc, f"Evidencia. {evidence}", bold_prefix="Evidencia. ")
    ui.add_body(doc, f"Para cerrar. {pending}", bold_prefix="Para cerrar. ")


def add_hu_next_steps_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [1450, 750, 3050, 4110]
    ui.set_table_geometry(table, widths)
    headers = ["HU", "Prio.", "Próxima acción", "Criterio de cierre"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        ui.set_cell_shading(cell, ui.INK)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT
        ui.set_font(p.add_run(text), size=9.1, color=ui.WHITE, bold=True)
    ui.set_repeat_table_header(table.rows[0])
    for hu_id, priority, action, close in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((hu_id, priority, action, close)):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = ui.Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT
            ui.set_font(p.add_run(text), size=8.8, bold=(idx in (0, 1)))
        ui.set_cell_shading(
            cells[1],
            ui.RED_LIGHT if priority == "P0" else ui.AMBER_LIGHT if priority == "P1" else ui.GRAY,
        )
    return table


def main():
    doc = Document()
    ui.configure_document(doc)
    bullet_id = ui.add_numbering(doc, "bullet")
    decimal_id = ui.add_numbering(doc, "decimal")

    try:
        commit = subprocess.check_output(
            ["git", "rev-parse", "--short", "HEAD"], cwd=ROOT, text=True
        ).strip()
    except Exception:
        commit = "no disponible"

    core = doc.core_properties
    core.title = "Seguimiento de Historias de Usuario — Fases 1 y 2 — Radar de Oportunidades"
    core.subject = "HU de Fase 1 y Fase 2, criterios de aceptación, evidencia y próximos pasos"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = ui.Pt(12)
    p.paragraph_format.space_after = ui.Pt(0)
    ui.set_font(p.add_run("SEGUIMIENTO FUNCIONAL BASADO EN HU"), size=9.5, color=ui.AMBER, bold=True)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = ui.Pt(6)
    ui.set_font(title.add_run("Seguimiento de Historias de Usuario"), size=28, color=ui.INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = ui.Pt(18)
    ui.set_font(
        subtitle.add_run("Fases 1 y 2 · criterios de aceptación, evidencia y próximos pasos"),
        size=13.5,
        color=ui.PURPLE,
        bold=True,
    )

    meta = doc.add_table(rows=2, cols=2)
    ui.set_table_geometry(meta, [4680, 4680], indent=0)
    metadata = [
        ("Objeto", "Responder el avance por historias de usuario"),
        ("Corte funcional", "25 de julio de 2026"),
        ("Versión de código", commit),
        ("Estado", "Documento para validación del cliente"),
    ]
    for idx, (label, value) in enumerate(metadata):
        cell = meta.cell(idx // 2, idx % 2)
        ui.set_cell_shading(cell, ui.GRAY)
        ui.set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": ui.MID_GRAY},
            bottom={"val": "single", "sz": "4", "color": ui.MID_GRAY},
            start={"val": "single", "sz": "4", "color": ui.MID_GRAY},
            end={"val": "single", "sz": "4", "color": ui.MID_GRAY},
        )
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = ui.Pt(2)
        ui.set_font(p1.add_run(label.upper()), size=8, color=ui.MUTED, bold=True)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = ui.Pt(0)
        ui.set_font(p2.add_run(value), size=10.2, color=ui.INK, bold=True)

    doc.add_paragraph()
    add_hu_metric_strip(doc)
    doc.add_paragraph()
    ui.add_callout(
        doc,
        "Respuesta directa",
        "Sí se ha trabajado tanto en Fase 1 como en Fase 2. La Fase 1 concentra el motor de datos "
        "y análisis definido por las HU entregadas; la Fase 2 reúne acceso, monetización, alertas, "
        "operación, experiencia de búsqueda y análisis de rentabilidad. Ambas fases tienen "
        "implementación verificable, aunque conservan criterios pendientes.",
        fill=ui.PURPLE_LIGHT,
        accent=ui.PURPLE,
    )
    ui.add_body(
        doc,
        "El 84% corresponde al cumplimiento ponderado de la Fase 1. El 75% de Fase 2 es un indicador "
        "técnico de sus siete HU identificadas; no equivale todavía a aceptación contractual. No se "
        "presenta un porcentaje global hasta que el cliente apruebe el peso relativo de cada fase.",
        italic=True,
        color=ui.MUTED,
    )

    doc.add_page_break()
    ui.add_heading(doc, "1. Qué se entiende por historia abordada", 1)
    ui.add_body(
        doc,
        "Una historia se considera abordada cuando existe una necesidad de usuario identificada, criterios "
        "de aceptación revisados, una implementación demostrable y evidencia técnica. Una HU solo se considera "
        "cerrada cuando también se han resuelto sus pendientes y decisiones de negocio."
    )
    for text in [
        "Cumplido: el criterio está implementado y puede demostrarse.",
        "Parcial: existe implementación, pero falta cobertura, trazabilidad o una validación.",
        "Pendiente: aún no existe evidencia suficiente para aceptar el criterio.",
        "Validación del cliente: decisión funcional que no debe asumirse desde desarrollo.",
    ]:
        ui.add_list_item(doc, text, bullet_id)

    ui.add_heading(doc, "2. Matrices maestras por fase", 1)
    ui.add_heading(doc, "2.1 Fase 1 — Motor de datos y análisis", 2)
    add_master_hu_table(
        doc,
        [
            ("HU-F1-01", "Comparar una propiedad con mercado equivalente y obtener clasificación CRECE.", "95%", "Avanzada", "Auditar histórico y aprobar regla geográfica."),
            ("HU-F1-02", "Conocer vigencia, cambios y retiro de publicaciones del Portal.", "78%", "Parcial", "Uniformar eventos del ciclo de vida."),
            ("HU-F1-03", "Analizar activos bancarios aunque la fuente omita atributos.", "82%", "Avanzada", "Cerrar identidad e historial entre bancos."),
            ("HU-F1-04", "Extraer y comparar activos AVAL desde sus boletines PDF.", "72%", "Parcial", "Normalizar ciudades y comparar boletines."),
            ("HU-F1-05", "Entender riesgo, cuota-parte y acceso de un remate judicial.", "92%", "Avanzada", "Cuarentena de anomalías y evidencia visual."),
        ],
    )
    ui.add_callout(
        doc,
        "Lectura de Fase 1",
        "Las HU del motor no están simplemente “iniciadas”: todas tienen implementación concreta. "
        "El trabajo restante consiste en cerrar criterios específicos, no en reconstruir el producto desde cero.",
        fill=ui.GOLD_LIGHT,
        accent=ui.AMBER,
    )
    doc.add_page_break()
    ui.add_heading(doc, "2.2 Fase 2 — Acceso, producto y operación", 2)
    add_complementary_hu_table(
        doc,
        [
            ("HU-F2-01", "Ingresar con Google y mantener una sesión segura.", "95%", "Validada manualmente", "Monitoreo y política de acceso."),
            ("HU-F2-02", "Acceder a funciones según cuenta y plan freemium/Pro.", "80%", "Demo funcional", "Permisos y activación real."),
            ("HU-F2-03", "Recibir alertas profesionales por correo.", "85%", "Operativa", "Agenda, métricas y dominio."),
            ("HU-F2-04", "Administrar usuarios, intereses y operación.", "70%", "Parcial", "Roles, auditoría y piloto."),
            ("HU-F2-05", "Comparar canon y rentabilidad de arriendo.", "68%", "Piloto", "Cobertura, muestra y confianza."),
            ("HU-F2-06", "Pagar un plan y activar beneficios Pro.", "35%", "Preparada", "Credenciales reales y webhook."),
            ("HU-F2-07", "Filtrar y navegar con claridad en desktop y móvil.", "95%", "Desplegada", "Prueba con usuarios y WCAG."),
        ],
    )
    ui.add_callout(
        doc,
        "Lectura de Fase 2",
        "La Fase 2 no está pendiente de iniciar: Google, alertas, planes, cuenta, administración, "
        "comparador, arriendos y filtros ya tienen evidencia. Wompi se mantiene deliberadamente "
        "incompleto hasta contar con datos legítimos del comercio.",
        fill=ui.PURPLE_LIGHT,
        accent=ui.PURPLE,
    )

    doc.add_page_break()
    ui.add_heading(doc, "3. Trazabilidad detallada de las HU de Fase 1", 1)
    add_user_story_block(
        doc,
        bullet_id,
        "HU-F1-01",
        "Núcleo de comparables e Índice CRECE",
        "95%",
        "Avanzada — próxima a cierre",
        "green",
        "Como inversionista inmobiliario, quiero comparar una propiedad con inmuebles equivalentes del mercado abierto para saber si su precio representa una oportunidad real.",
        [
            ("done", "Calcular mediana robusta por precio por metro cuadrado y recortar valores atípicos."),
            ("done", "Comparar por tipo, área, atributos y cercanía geográfica."),
            ("done", "Excluir bancos y remates del universo de referencia del mercado abierto."),
            ("done", "Traducir el resultado a las once categorías del Índice CRECE."),
            ("partial", "Persistir y auditar el índice para todas las filas históricas."),
        ],
        "Motor de comparables e Índice CRECE, criterios visibles para el usuario y pruebas sobre límites de clasificación.",
        "Completar auditoría histórica y formalizar el radio o regla de ampliación geográfica.",
        "Aprobar cuándo se amplía de barrio a zona o ciudad y cuál es el radio máximo aceptable.",
    )

    doc.add_page_break()
    add_user_story_block(
        doc,
        bullet_id,
        "HU-F1-02",
        "Rango de precio y vigencia del Portal",
        "78%",
        "Parcial — reglas listas, historial incompleto",
        "amber",
        "Como usuario del Radar, quiero saber si una publicación es nueva, vigente, modificada o retirada para no analizar inventario desactualizado.",
        [
            ("done", "Configurar ciudades, rangos de precio, área, estrato y fecha de revisión."),
            ("done", "Aplicar reglas de vigencia de 30 y 90 días."),
            ("done", "Modelar eventos de vigencia y una cadencia semanal."),
            ("partial", "Escribir todos los eventos desde cada scraper de forma uniforme."),
            ("partial", "Verificar automáticamente la activación y población de ciudades."),
        ],
        "Configuración persistente de zonas, motor de vigencia y pruebas sobre estados temporales.",
        "Cerrar NUEVO, VISTO, PRECIO_ACTUALIZADO, RETIRADO, EXPIRADO y DESCARTADO en todas las fuentes.",
        "Definir si una publicación retirada debe conservarse como histórico visible y durante cuánto tiempo.",
    )

    doc.add_page_break()
    add_user_story_block(
        doc,
        bullet_id,
        "HU-F1-03",
        "Motor de activos bancarios",
        "82%",
        "Avanzada — cierre de identidad e historial",
        "amber",
        "Como inversionista, quiero comparar activos bancarios contra el mercado abierto aunque el banco no publique todos los atributos, para detectar oportunidades sin descartar inventario útil.",
        [
            ("done", "Normalizar fuentes bancarias y compararlas contra mercado abierto."),
            ("done", "Tratar el estrato ausente como dato opcional."),
            ("done", "Mostrar la última verificación y rotar el inventario semanalmente."),
            ("done", "Adaptar filtros e interfaz a la calidad real de los datos bancarios."),
            ("partial", "Resolver identidad secundaria, retiro, cambio de precio e historial entre bancos."),
        ],
        "Scrapers bancarios, clasificación CRECE, rotación semanal y pruebas de frescura.",
        "Cerrar la estrategia de identidad para enlaces o slugs inestables y unificar eventos.",
        "Confirmar qué atributos mínimos permiten publicar un activo cuando faltan estrato, habitaciones o baños.",
    )

    doc.add_page_break()
    add_user_story_block(
        doc,
        bullet_id,
        "HU-F1-04",
        "Motor de activos Grupo AVAL",
        "72%",
        "Parcial — ingesta funcional, normalización pendiente",
        "amber",
        "Como usuario, quiero consultar activos AVAL extraídos de sus boletines PDF y conocer sus cambios para analizarlos junto con las demás oportunidades.",
        [
            ("done", "Resolver, descargar y leer dinámicamente el boletín PDF."),
            ("done", "Deduplicar por código y extraer precio, tipo, área, dirección y ciudad."),
            ("done", "Evitar comparaciones inválidas mediante area_tipo."),
            ("partial", "Normalizar ciudad–departamento y marcar datos que requieren revisión."),
            ("partial", "Comparar boletines e identificar de forma trazable novedades y cambios de valor."),
        ],
        "Parser PDF, deduplicación, modelo normalizado e integración de las fichas en la aplicación.",
        "Completar normalización, comparación entre boletines, campos opcionales y trazabilidad documental.",
        "Aprobar si las insignias “Nuevo inmueble” y “Nuevo valor” son suficientes o deben recalcularse.",
    )

    doc.add_page_break()
    add_user_story_block(
        doc,
        bullet_id,
        "HU-F1-05",
        "Motor de remates judiciales",
        "92%",
        "Avanzada — cierre de calidad de datos",
        "green",
        "Como inversionista, quiero distinguir el riesgo jurídico, la postura mínima y el porcentaje real del bien rematado para evitar interpretar una subasta de forma equivocada.",
        [
            ("done", "Identificar origen del demandante, tipo de proceso y cuota-parte."),
            ("done", "Aplicar una regla conservadora cuando el origen no es claro."),
            ("done", "Diferenciar porcentaje de postura y porcentaje del dominio rematado."),
            ("done", "Mostrar la matriz de acceso y alertas jurídicas."),
            ("partial", "Enviar anomalías a cuarentena y usar evidencia neutral cuando no exista imagen."),
        ],
        "Motor jurídico, matriz comercial, sanitización de datos y pruebas específicas sobre avisos y porcentajes.",
        "Implementar cuarentena desde ingestión/base y cerrar la evidencia visual trazable.",
        "Definir qué evidencia visual o documental permite publicar una ficha sin imagen original.",
    )

    doc.add_page_break()
    ui.add_heading(doc, "4. Trazabilidad detallada de las HU de Fase 2", 1)
    ui.add_body(
        doc,
        "Estas HU traducen la hoja de ruta del producto y las decisiones posteriores del cliente a "
        "criterios verificables. El porcentaje de Fase 2 es un avance técnico provisional hasta que "
        "el cliente apruebe formalmente esta línea base."
    )
    add_phase2_story_block(
        doc,
        bullet_id,
        "HU-F2-01",
        "Inicio de sesión con Google",
        "95%",
        "Validada — operativa",
        "green",
        "Como usuario, quiero ingresar con mi cuenta de Google para acceder al Radar sin crear otra contraseña.",
        [
            ("done", "Mostrar el acceso con Google en la pantalla de ingreso."),
            ("done", "Completar OAuth y conservar la sesión del usuario."),
            ("partial", "Monitorear altas, errores y políticas de acceso en operación."),
        ],
        "Proveedor OAuth y retorno validados manualmente; el E2E local comprueba disponibilidad del acceso y reglas de autenticación sin automatizar la pantalla externa de Google.",
        "Acordar política de acceso y monitorear los primeros usuarios reales.",
    )
    add_phase2_story_block(
        doc,
        bullet_id,
        "HU-F2-02",
        "Cuenta, planes y acceso freemium/Pro",
        "80%",
        "Demo funcional — permisos pendientes",
        "amber",
        "Como usuario, quiero conocer mi plan y acceder a beneficios acordes con mi suscripción para entender el valor de pasar a Pro.",
        [
            ("done", "Mostrar cuenta, planes y beneficios disponibles."),
            ("done", "Presentar contenido restringido y llamadas a suscripción."),
            ("partial", "Aplicar permisos de forma consistente en servidor y base de datos."),
        ],
        "Vistas de cuenta y planes, muro freemium, intereses de plan y checkout de demostración.",
        "Validar entitlements, políticas RLS y activación/cancelación del plan.",
    )

    doc.add_page_break()
    add_phase2_story_block(
        doc,
        bullet_id,
        "HU-F2-03",
        "Alertas por correo",
        "85%",
        "Operativa — endurecimiento pendiente",
        "green",
        "Como usuario, quiero recibir alertas por correo cuando aparezcan oportunidades compatibles para reaccionar a tiempo.",
        [
            ("done", "Crear alertas asociadas a preferencias del usuario."),
            ("done", "Enviar un correo profesional mediante Resend."),
            ("partial", "Ejecutar la agenda con idempotencia, métricas y control de rebotes."),
        ],
        "Despacho de alertas, plantilla brandeada y prueba real de correo.",
        "Verificar dominio, activar agenda controlada y medir entrega, rebotes y duplicados.",
    )
    add_phase2_story_block(
        doc,
        bullet_id,
        "HU-F2-04",
        "Administración y operación",
        "70%",
        "Parcial — apta para piloto interno",
        "amber",
        "Como administrador, quiero revisar usuarios, intereses y estado operativo para controlar el servicio.",
        [
            ("done", "Consultar información de cuenta e intereses de plan."),
            ("partial", "Restringir las funciones administrativas por rol."),
            ("partial", "Registrar auditoría de cambios y eventos operativos."),
        ],
        "Servicios de cuenta, administración comercial, estadísticas y logs de scraping.",
        "Cerrar roles, auditoría, indicadores operativos y prueba piloto con un administrador.",
    )

    doc.add_page_break()
    add_phase2_story_block(
        doc,
        bullet_id,
        "HU-F2-05",
        "Comparables de arriendo y rentabilidad",
        "68%",
        "Piloto — cobertura por consolidar",
        "amber",
        "Como inversionista, quiero estimar el canon de arriendo de una propiedad y su rentabilidad con comparables similares para tomar una decisión más completa.",
        [
            ("done", "Separar ofertas de venta y arriendo y almacenar el canon."),
            ("done", "Calcular comparables de alquiler y una estimación de rentabilidad."),
            ("partial", "Mostrar cobertura, muestra mínima y nivel de confianza."),
        ],
        "Piloto para Medellín, Bogotá, Cali, Barranquilla y Bucaramanga; recorrido E2E de rentabilidad aprobado.",
        "Ejecutar ingestas recurrentes, medir cobertura y aprobar reglas de muestra y confianza.",
    )
    add_phase2_story_block(
        doc,
        bullet_id,
        "HU-F2-06",
        "Pago y activación de plan",
        "35%",
        "Preparada — bloqueada por datos del comercio",
        "amber",
        "Como usuario, quiero pagar un plan de forma segura para activar los beneficios Pro.",
        [
            ("done", "Preparar la experiencia de checkout y el flujo de demostración."),
            ("partial", "Integrar Wompi Sandbox con credenciales legítimas."),
            ("pending", "Validar webhook, conciliación y activación idempotente del plan."),
        ],
        "Pantalla y flujo de checkout de demostración listos.",
        "Retomar únicamente cuando existan datos reales del comercio; completar transacción Sandbox y webhook firmado.",
    )

    doc.add_page_break()
    add_phase2_story_block(
        doc,
        bullet_id,
        "HU-F2-07",
        "Filtros y navegación responsive",
        "95%",
        "Desplegada — validación con usuarios pendiente",
        "green",
        "Como usuario, quiero filtrar oportunidades con controles claros en desktop y móvil para encontrar propiedades con menos esfuerzo.",
        [
            ("done", "Usar un panel lateral de filtros en desktop aprovechando el espacio disponible."),
            ("done", "Adaptar los filtros y resultados a móvil sin perder el contexto."),
            ("partial", "Validar accesibilidad y comprensión con usuarios representativos."),
        ],
        "Panel lateral ampliado, filtros responsive y recorridos E2E de búsqueda aprobados.",
        "Ejecutar una sesión de usabilidad y cerrar hallazgos WCAG prioritarios.",
    )
    ui.add_callout(
        doc,
        "Alcance transversal de Fase 2",
        "El despliegue público, la seguridad de credenciales, la observabilidad y la robustez de producción "
        "son habilitadores técnicos de estas HU. Se reportan como condiciones de salida, no como historias "
        "independientes, para evitar duplicar el porcentaje.",
        fill=ui.GOLD_LIGHT,
        accent=ui.AMBER,
    )

    doc.add_page_break()
    ui.add_heading(doc, "5. Próximos pasos vinculados a las HU", 1)
    add_hu_next_steps_table(
        doc,
        [
            ("Todas", "P0", "Validar esta matriz y convertirla en backlog oficial.", "Cada HU tiene redacción, criterios, prioridad y aprobación."),
            ("HU-F1-02/03", "P0", "Uniformar eventos de ciclo de vida Portal y Bancos.", "Las muestras tienen historial completo y reproducible."),
            ("HU-F1-04", "P1", "Cerrar normalización y comparación de boletines AVAL.", "Ciudades, duplicados y cambios quedan trazables."),
            ("HU-F2-05", "P1", "Consolidar piloto de arriendos.", "El canon solo aparece con cobertura y confianza acordadas."),
            ("HU-F2-03", "P1", "Operar alertas recurrentes con canary.", "Cada alerta llega una vez y queda registrada."),
            ("HU-F2-02/04", "P1", "Cerrar permisos, roles y auditoría.", "Cuenta, plan y administración respetan políticas verificables."),
            ("HU-F2-06", "P2", "Retomar Wompi cuando existan credenciales.", "Una transacción Sandbox aprobada activa Pro correctamente."),
        ],
    )

    ui.add_heading(doc, "6. Preguntas de validación por HU", 1)
    for text in [
        "HU-F1-01: ¿qué radio o ampliación geográfica debe aceptar el motor de comparables?",
        "HU-F1-02: ¿durante cuánto tiempo debe mostrarse el histórico de retirados o expirados?",
        "HU-F1-03/HU-F1-04: ¿qué atributos mínimos permiten publicar una ficha incompleta?",
        "HU-F1-04: ¿las insignias del boletín AVAL son autoritativas o deben recalcularse?",
        "HU-F1-05: ¿qué evidencia es suficiente cuando un remate no tiene imagen original?",
        "HU-F2-02: ¿qué funcionalidades exactas pertenecen a cada plan y qué límites aplican?",
        "HU-F2-03: ¿con qué frecuencia y bajo qué preferencias deben enviarse alertas?",
        "HU-F2-05: ¿qué ciudades, tipos de inmueble y muestra mínima se priorizan para arriendos?",
        "HU-F2-06: ¿se confirma que Wompi queda aplazado hasta contar con datos reales del comercio?",
        "Fases 1 y 2: ¿qué peso tendrá cada fase para calcular un porcentaje global aceptado?",
    ]:
        ui.add_list_item(doc, text, bullet_id)

    doc.add_page_break()
    ui.add_heading(doc, "7. Guion para presentar el avance por historias", 1)
    for text in [
        "Explicar que se trazaron cinco HU de Fase 1 y siete HU de Fase 2.",
        "Presentar por separado el 84% de Fase 1 y el 75% técnico de Fase 2.",
        "Demostrar HU-F1-01 con una propiedad y sus comparables CRECE.",
        "Demostrar HU-F1-03 y HU-F1-05 con un activo bancario y un remate.",
        "Demostrar HU-F2-01, HU-F2-03 y HU-F2-07 con acceso, alertas y filtros.",
        "Mostrar HU-F2-05 como piloto y HU-F2-06 como integración preparada, no terminada.",
        "Revisar los criterios parciales sin llamarlos terminados.",
        "Cerrar solicitando respuesta a las preguntas de la sección 6.",
    ]:
        ui.add_list_item(doc, text, decimal_id)
    ui.add_callout(
        doc,
        "Mensaje central",
        "Las historias de Fase 1 y Fase 2 se están ejecutando y cuentan con evidencia. "
        "El siguiente paso es validar criterios, aprobar el peso de cada fase y convertir esta matriz "
        "en el control oficial de alcance y aceptación.",
        fill=ui.PURPLE_LIGHT,
        accent=ui.PURPLE,
    )

    ui.add_heading(doc, "8. Evidencia técnica asociada", 1)
    for text in [
        "TypeScript: compilación sin errores.",
        "Pruebas unitarias e integración: suite completa aprobada sin fallos en el corte.",
        "Recorridos E2E locales: suite crítica aprobada sin fallos en el corte.",
        "E2E cubiertos: dashboard, Bancos, Remates, arriendos/rentabilidad, guardados, login, disponibilidad del acceso con Google, checkout demo, móvil, filtros y estados de carga.",
        f"Versión evaluada: {commit}.",
        "Producción: https://joinsclee-radar.juno8i.easypanel.host/",
    ]:
        ui.add_list_item(doc, text, bullet_id)

    doc.add_page_break()
    ui.add_heading(doc, "9. Documentos fuente de las HU", 1)
    for text in [
        "Radar de Oportunidades — Documentación y Propuesta.",
        "Radar CRECE — Índice Maestro HU Fase 1.",
        "Radar CRECE — Especificación Motor de Comparables v1.2.",
        "HU Núcleo Motor de Comparables.",
        "HU Rango de Precio y Vigencia del Motor Radar.",
        "HU Motor de Activos Bancarios.",
        "HU Motor de Activos Bancarios — Grupo AVAL.",
        "HU Motor de Remates Judiciales.",
        "Hoja de ruta del documento inicial: Google, freemium, alertas, panel, comparador y rentabilidad.",
        "Decisiones posteriores del cliente: Resend, Wompi, filtros laterales y comparables de arriendo.",
    ]:
        ui.add_list_item(doc, text, bullet_id)
    ui.add_callout(
        doc,
        "Control recomendado",
        "A partir de la validación, ninguna funcionalidad nueva debería entrar al alcance sin un ID de HU, "
        "criterios de aceptación, prioridad y aprobación del cliente.",
        fill=ui.GOLD_LIGHT,
        accent=ui.AMBER,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
