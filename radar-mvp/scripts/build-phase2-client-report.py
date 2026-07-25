#!/usr/bin/env python3
"""Genera el informe verificable de Fase 2 en DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "docs" / "Informe_Avance_Radar_Fase2_2026-07-24.docx"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "202124"
MUTED = "5F6368"
LINE = "DADCE0"
LIGHT = "F5F7FA"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "EAF5EF"
PALE_AMBER = "FFF6DD"
GREEN = "237A57"
AMBER = "8A5B00"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360
TABLE_INDENT = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_font(run, size=11, color=INK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:color"), color)


def ensure(parent, tag):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def apply_table_geometry(table, widths):
    if sum(widths) != TABLE_WIDTH:
        raise ValueError(f"Las columnas deben sumar {TABLE_WIDTH}: {widths}")
    table.autofit = False
    table.alignment = 0
    tbl_pr = table._tbl.tblPr
    tbl_w = ensure(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_ind = ensure(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    layout = ensure(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)

    for index, width in enumerate(widths):
        table.columns[index].width = Twips(width)
    for row in table.rows:
        row.height = None
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Twips(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = ensure(tc_pr, "w:tcW")
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            margins = ensure(tc_pr, "w:tcMar")
            for side, value in CELL_MARGINS.items():
                margin = ensure(margins, f"w:{side}")
                margin.set(qn("w:type"), "dxa")
                margin.set(qn("w:w"), str(value))


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_font(run, size=8, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, end])


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(3)
    left = hp.add_run("RADAR DE OPORTUNIDADES  /  FASE 2")
    set_font(left, size=8, color=NAVY, bold=True)
    right = hp.add_run("     INFORME VERIFICABLE")
    set_font(right, size=8, color=BLUE, bold=True)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), LINE)
    p_bdr.append(bottom)
    hp._p.get_or_add_pPr().append(p_bdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Estado local validado · No representa activación de cobros o correo"
    set_font(fp.runs[0], size=8, color=MUTED)
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    footer_table.cell(0, 0).paragraphs[0].add_run("Radar CRECE · 24 jul 2026")
    set_font(footer_table.cell(0, 0).paragraphs[0].runs[0], size=8, color=MUTED)
    page_number(footer_table.cell(0, 1).paragraphs[0])
    apply_table_geometry(footer_table, [7800, 1560])
    footer._element.remove(fp._element)

    core = doc.core_properties
    core.title = "Radar de Oportunidades - Informe de avance Fase 2"
    core.subject = "Estado funcional, QA, stack y pendientes externos"
    core.author = "Equipo Radar CRECE"
    core.keywords = "Radar CRECE, Fase 2, planes, alertas, comparador, rentabilidad"


def title_block(doc):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(18)
    kicker.paragraph_format.space_after = Pt(6)
    run = kicker.add_run("INFORME DE AVANCE DE PRODUCTO")
    set_font(run, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Radar de Oportunidades\nFase 2")
    set_font(run, size=26, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Base comercial, alertas, comparador, rentabilidad y operación")
    set_font(run, size=13, color=MUTED)

    callout(doc, "Conclusión", "Fase 1 permanece en 84%. La Fase 2 comercial alcanza 67% con un recorrido funcional validado en local; pagos, correo entregado, canon observado e integraciones todavía requieren decisiones externas.", "67%", PALE_BLUE, NAVY)

    meta = [
        ("Preparado para", "Andrés Giraldo"),
        ("Corte", "24 de julio de 2026"),
        ("Estado", "Listo para demostración local; no desplegado"),
        ("Versión local", "f606520"),
        ("Producción vigente", "joinsclee-radar.juno8i.easypanel.host"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    for index, (label, value) in enumerate(meta):
        left, right = table.rows[index].cells
        set_cell_fill(left, LIGHT)
        set_cell_border(left)
        set_cell_border(right)
        p = left.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(label.upper()), size=8, color=BLUE, bold=True)
        p = right.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(value), size=10, color=INK, bold=True)
    apply_table_geometry(table, [2300, 7060])

    doc.add_paragraph()
    callout(doc, "Lectura honesta", "El término Fase 2 se usó históricamente para remates y después para la expansión comercial. Este informe separa ambos alcances para no inflar el avance.", None, PALE_AMBER, AMBER)


def callout(doc, title, body, value=None, fill=PALE_BLUE, accent=BLUE):
    widths = [1800, 7560] if value else [9360]
    table = doc.add_table(rows=1, cols=len(widths))
    if value:
        value_cell, text_cell = table.rows[0].cells
        set_cell_fill(value_cell, accent)
        p = value_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(value), size=24, color=WHITE, bold=True)
    else:
        text_cell = table.cell(0, 0)
    set_cell_fill(text_cell, fill)
    p = text_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(title), size=11, color=accent, bold=True)
    p = text_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    set_font(p.add_run(body), size=10, color=INK)
    for cell in table.rows[0].cells:
        set_cell_border(cell, fill)
    apply_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), bold=True)
        set_font(p.add_run(text[len(bold_lead):]))
    else:
        set_font(p.add_run(text))
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_font(p.add_run(text))
    return p


def new_numbering_sequence(doc, style_name="List Number"):
    numbering = doc.part.numbering_part.element
    style_num_id = doc.styles[style_name].element.pPr.numPr.numId.val
    source_num = next(
        num for num in numbering.findall(qn("w:num"))
        if int(num.get(qn("w:numId"))) == int(style_num_id)
    )
    abstract_num_id = source_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    next_num_id = max(
        [int(num.get(qn("w:numId"))) for num in numbering.findall(qn("w:num"))],
        default=0,
    ) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return next_num_id


def numbered(doc, text, num_id=None):
    p = doc.add_paragraph(style="List Number")
    if num_id is not None:
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = num_id
    set_font(p.add_run(text))
    return p


def table(doc, headers, rows, widths, font_size=9):
    result = doc.add_table(rows=1, cols=len(headers))
    repeat_header(result.rows[0])
    for index, (cell, label) in enumerate(zip(result.rows[0].cells, headers)):
        set_cell_fill(cell, NAVY)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(label), size=9, color=WHITE, bold=True)
    for row_index, values in enumerate(rows):
        cells = result.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_fill(cell, WHITE if row_index % 2 == 0 else LIGHT)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_font(p.add_run(str(value)), size=font_size)
    apply_table_geometry(result, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return result


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build():
    doc = Document()
    configure(doc)
    title_block(doc)

    page_break(doc)
    heading(doc, "1. Resumen ejecutivo")
    body(doc, "El compromiso de llegar a por lo menos 80% de la Fase 1 ya se cumplió. El corte defendible de esa fase sigue siendo 84%. La ejecución actual agrega una base comercial y operativa de Fase 2 sin convertir maquetas en promesas.")
    heading(doc, "Dos significados de Fase 2", 2)
    bullet(doc, "Fase 2 técnica original: motor de remates judiciales, aproximadamente 92%.")
    bullet(doc, "Fase 2 comercial/expansión: planes, alertas, rentabilidad, comparador, administración, pagos e integraciones, aproximadamente 67%.")
    callout(doc, "Dictamen para el cliente", "El producto supera el 80% comprometido para Fase 1 y la Fase 2 comercial ya tiene un recorrido demostrable en local. Todavía no debe venderse como monetización o notificaciones plenamente operativas.", None, PALE_GREEN, GREEN)

    heading(doc, "Qué puede demostrarse hoy", 2)
    for item in [
        "Planes Free y Pro con límites de acceso reales.",
        "Centro de cuenta, sincronización y alertas persistentes.",
        "Comparador de dos o tres inmuebles guardados.",
        "Rentabilidad bruta y neta a partir de canon declarado.",
        "Panel administrativo agregado y exportación de cuenta.",
        "Degradación segura cuando correo o pagos no están configurados.",
    ]:
        bullet(doc, item)

    page_break(doc)
    heading(doc, "2. Entregables del corte")
    sections = [
        ("Planes y acceso", [
            "Catálogo API y página pública /planes.",
            "Precio Pro explícitamente por definir; no hay cobro simulado.",
            "Solicitud de plan y compatibilidad con marcas históricas de suscripción.",
            "Una alerta para Free y hasta cinco para Pro.",
        ]),
        ("Cuenta y alertas", [
            "Centro /cuenta y continuidad de preferencias, simulaciones y alertas.",
            "Validación separada de entradas y alertas persistidas.",
            "Historial de entrega, idempotencia y reintentos escalonados.",
            "Adaptador Resend y proceso semanal protegido.",
        ]),
        ("Comparación y rentabilidad", [
            "Comparador /comparador con criterios homogéneos.",
            "Respeto del muro de acceso al recuperar favoritos.",
            "Rentabilidad bruta y neta con vacancia, mantenimiento y administración.",
            "Exportación CSV y separación entre canon declarado y observado.",
        ]),
        ("Administración", [
            "Panel /admin protegido por rol.",
            "Usuarios, embudo Pro, alertas, perfiles y estado de entregas.",
            "Exportación JSON sin contraseñas ni tokens.",
        ]),
    ]
    for title, items in sections:
        heading(doc, title, 2)
        for item in items:
            bullet(doc, item)

    page_break(doc)
    heading(doc, "3. Avance ponderado de Fase 2 comercial")
    score_rows = [
        ["Base comercial, cuenta y UX", "15%", "90%", "13,5"],
        ["Planes, acceso y suscripción", "15%", "65%", "9,8"],
        ["Alertas y notificaciones", "20%", "85%", "17,0"],
        ["Comparador y exportación", "15%", "85%", "12,8"],
        ["Canon y rentabilidad", "15%", "35%", "5,3"],
        ["Administración", "10%", "75%", "7,5"],
        ["Pagos e integraciones", "10%", "10%", "1,0"],
        ["TOTAL", "100%", "", "66,8% ≈ 67%"],
    ]
    table(doc, ["Bloque", "Peso", "Cumpl.", "Aporte"], score_rows, [5000, 1200, 1500, 1660], 8.7)
    body(doc, "Este porcentaje mide alcance adicional y no reduce el 84% de Fase 1. La diferencia hasta 80% se concentra en proveedores, datos autorizados y operación comercial.")
    heading(doc, "Evidencia de calidad local", 2)
    evidence = [
        ["TypeScript", "Sin errores"],
        ["Unitarias/integración", "94/94"],
        ["Recorridos E2E", "6/6"],
        ["Errores de consola", "0"],
        ["QA visual", "1440×1000 y 375×812"],
        ["Salud ponderada", "98,6/100"],
        ["Usuarios reales alterados", "0"],
    ]
    table(doc, ["Control", "Resultado"], evidence, [4800, 4560], 9)

    page_break(doc)
    heading(doc, "4. Stack y arquitectura")
    stack = [
        ["Runtime", "Node.js 20+, TypeScript, ESM"],
        ["Servidor", "HTTP nativo de Node y API JSON"],
        ["Frontend", "HTML, CSS y JavaScript nativos"],
        ["Datos", "Supabase PostgreSQL, Auth y Storage"],
        ["Validación", "Zod"],
        ["Extracción", "Firecrawl, Playwright y parsers PDF"],
        ["Motor", "Comparables, estadística robusta e Índice CRECE"],
        ["Automatización", "node-cron y scheduler persistido"],
        ["QA", "node:test, Playwright y Chromium"],
        ["Infraestructura", "Docker, VPS y EasyPanel"],
    ]
    table(doc, ["Capa", "Tecnología / decisión"], stack, [2600, 6760], 9)
    callout(doc, "Arquitectura comercial", "La cuenta y las alertas reutilizan Supabase Auth para el corte inicial. Cuando el volumen o la analítica lo exijan, deben migrarse a tablas dedicadas con historial de entregas.", None, PALE_BLUE, BLUE)
    heading(doc, "Controles incorporados", 2)
    for item in [
        "Autorización Bearer en cuenta, favoritos y administración.",
        "Rol administrativo explícito.",
        "Contenido premium redactado antes de salir del servidor.",
        "Escape de contenido HTML en correos.",
        "Idempotencia de proveedor, historial y reintentos.",
        "Límites de abuso para endpoints sensibles.",
        "Proceso interno cerrado si falta secreto o proveedor.",
        "Límites y validación de payloads.",
    ]:
        bullet(doc, item)

    page_break(doc)
    heading(doc, "5. Ruta de 67% a 80%")
    heading(doc, "Prioridad 1: activar alertas", 2)
    for item in [
        "Aprobar Resend, verificar dominio y configurar las cuatro variables.",
        "Crear el trabajo semanal en EasyPanel.",
        "Ejecutar un envío controlado y verificar historial, deduplicación y reintentos.",
    ]:
        numbered(doc, item)
    heading(doc, "Prioridad 2: cerrar planes y pago", 2)
    for item in [
        "Aprobar precio, moneda, periodicidad, prueba y cancelación.",
        "Elegir proveedor de pago.",
        "Implementar checkout, webhooks, renovación y fallo de cobro.",
    ]:
        numbered(doc, item)
    heading(doc, "Prioridad 3: canon observado", 2)
    for item in [
        "Definir una fuente legal/autorizada.",
        "Normalizar y trazar vigencia por ciudad, zona, tipo y área.",
        "Separar rentabilidad declarada de rentabilidad basada en datos.",
    ]:
        numbered(doc, item)
    callout(doc, "Condición para superar 80%", "Alertas entregadas, checkout funcional y una primera fuente de canon autorizada. Las integraciones Low Ticket/Tradentia pueden quedar después si el tercero no entrega especificación.", None, PALE_GREEN, GREEN)

    page_break(doc)
    heading(doc, "6. Dependencias externas")
    dependencies = [
        ["Precio Radar Pro", "Tarifa y beneficios", "Página y catálogo listos"],
        ["Proveedor de pago", "Elegir proveedor", "No seleccionado"],
        ["Resend", "Cuenta, dominio y API key", "Adaptador listo"],
        ["EasyPanel cron", "Secreto y horario", "Endpoint listo"],
        ["Canon", "Fuente autorizada", "Cálculo manual listo"],
        ["Low Ticket/Tradentia", "API o especificación", "Pendiente externo"],
        ["Administrador", "Usuario autorizado", "Control listo"],
    ]
    table(doc, ["Dependencia", "Decisión", "Código"], dependencies, [2500, 3600, 3260], 8.6)
    heading(doc, "Variables para alertas", 2)
    for item in [
        "RESEND_API_KEY",
        "ALERTS_FROM_EMAIL",
        "ALERTS_CRON_SECRET",
        "APP_BASE_URL",
    ]:
        bullet(doc, item)
    callout(doc, "Comportamiento seguro", "Sin estas variables la alerta se guarda, pero la aplicación informa que el correo sigue pendiente. El proceso devuelve 503 y no marca entregas falsas.", None, PALE_AMBER, AMBER)

    page_break(doc)
    heading(doc, "7. Límites y endurecimiento")
    heading(doc, "No debe afirmarse todavía", 2)
    for item in [
        "Que Radar Pro ya cobra o renueva.",
        "Que las alertas ya llegan a correos reales.",
        "Que el canon proviene de una fuente de mercado.",
        "Que Low Ticket o Tradentia ya están integrados.",
        "Que existe un SLA empresarial.",
    ]:
        bullet(doc, item)
    heading(doc, "Antes de apertura comercial", 2)
    hardening = [
        ["Sesión", "Cookies HttpOnly/Secure/SameSite y CSRF"],
        ["Abuso", "Contadores compartidos al operar varias réplicas"],
        ["Identidad", "Verificación de correo"],
        ["Operación", "CI/CD, métricas, alarmas y trazas"],
        ["Recuperación", "Restore de backup y rollback"],
        ["Calidad", "Carga, canary y auditoría WCAG"],
    ]
    table(doc, ["Frente", "Resultado esperado"], hardening, [2300, 7060], 9)

    page_break(doc)
    heading(doc, "8. Guion de demostración")
    demo = [
        "Abrir /planes y explicar por qué no se publica una tarifa no aprobada.",
        "Iniciar sesión y abrir /cuenta.",
        "Mostrar preferencias sincronizadas y crear una alerta.",
        "Guardar dos inmuebles y abrir /comparador.",
        "Ingresar canon y administración en una ficha.",
        "Descargar la exportación de cuenta.",
        "Abrir /admin solo con una cuenta autorizada.",
        "Cerrar con las dependencias para llegar a 80% de Fase 2.",
    ]
    demo_num_id = new_numbering_sequence(doc)
    for item in demo:
        numbered(doc, item, demo_num_id)
    callout(doc, "Frase de cierre", "“La Fase 1 ya superó el 80%. En Fase 2 ya existen planes, cuenta, alertas trazables, comparador, exportación y rentabilidad. Lo pendiente son decisiones y servicios externos que no debemos fingir.”", None, PALE_BLUE, NAVY)
    heading(doc, "Criterio de publicación", 2)
    body(doc, "No desplegar este corte hasta aprobar el texto comercial, configurar el primer canal externo y repetir el smoke test local. Después puede publicarse mediante PR y despliegue controlado.")

    return doc


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
