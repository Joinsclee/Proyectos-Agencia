#!/usr/bin/env python3
"""Genera el informe ejecutivo de avance de Radar CRECE en formato DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "docs" / "Informe_Avance_Radar_Oportunidades_Fase1_2026-07-24.docx"

NAVY = "12355B"
BLUE = "1F5A94"
MID_BLUE = "3D78A8"
PALE_BLUE = "E9F2FA"
ICE = "F5F8FB"
INK = "1F2937"
MUTED = "5B6775"
LINE = "CBD6E2"
GREEN = "237A57"
PALE_GREEN = "E8F4EE"
AMBER = "A96700"
PALE_AMBER = "FFF3D6"
RED = "A53A38"
PALE_RED = "FBE9E7"
WHITE = "FFFFFF"
PURPLE = "613174"
GOLD = "F2CA04"

PAGE_WIDTH_IN = 8.5
CONTENT_WIDTH_IN = 7.0


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        for key, value in edge_data.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_fixed(table) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_repeat_headers_false(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if keep is not None:
        p_pr.remove(keep)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, fld_end])


def set_run(run, size=None, color=INK, bold=False, italic=False, font="Arial") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_text(paragraph, text, *, bold=False, color=INK, size=None, italic=False):
    run = paragraph.add_run(text)
    set_run(run, size=size, color=color, bold=bold, italic=italic)
    return run


def add_callout(doc, title: str, body: str, *, fill=PALE_BLUE, accent=BLUE, value=None):
    table = doc.add_table(rows=1, cols=2 if value else 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    if value:
        widths = [1.55, CONTENT_WIDTH_IN - 1.55]
        table.columns[0].width = Inches(widths[0])
        table.columns[1].width = Inches(widths[1])
        value_cell, text_cell = table.rows[0].cells
        value_cell.width = Inches(widths[0])
        text_cell.width = Inches(widths[1])
        set_cell_shading(value_cell, accent)
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        vp = value_cell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(vp, value, bold=True, color=WHITE, size=26)
    else:
        text_cell = table.cell(0, 0)
        text_cell.width = Inches(CONTENT_WIDTH_IN)
    set_cell_shading(text_cell, fill)
    set_cell_margins(text_cell, top=140, bottom=140, start=170, end=170)
    set_cell_border(
        text_cell,
        left={"val": "single", "sz": "20", "color": accent},
        top={"val": "single", "sz": "2", "color": fill},
        bottom={"val": "single", "sz": "2", "color": fill},
        right={"val": "single", "sz": "2", "color": fill},
    )
    p = text_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_text(p, title, bold=True, color=accent, size=10.5)
    p2 = text_cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.05
    add_text(p2, body, color=INK, size=9.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_section_heading(doc, number: str, title: str, subtitle: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    add_text(p, number, bold=True, color=MID_BLUE, size=10)
    p.add_run("  ")
    add_text(p, title, bold=True, color=NAVY, size=19)
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(12)
        sp.paragraph_format.keep_with_next = True
        add_text(sp, subtitle, color=MUTED, size=9.5)
    else:
        p.paragraph_format.space_after = Pt(11)


def add_subheading(doc, title: str, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    add_text(p, title, bold=True, color=color, size=11.5)
    return p


def add_body(doc, text: str, *, bold_lead: str | None = None, color=INK, size=9.7):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.08
    if bold_lead and text.startswith(bold_lead):
        add_text(p, bold_lead, bold=True, color=color, size=size)
        add_text(p, text[len(bold_lead):], color=color, size=size)
    else:
        add_text(p, text, color=color, size=size)
    return p


def add_bullet(doc, text: str, *, color=INK, marker_color=BLUE, size=9.4, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18 + level * 0.2)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    add_text(p, "• ", bold=True, color=marker_color, size=size)
    add_text(p, text, color=color, size=size)
    return p


def add_numbered(doc, number: int, title: str, body: str):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    table.columns[0].width = Inches(0.48)
    table.columns[1].width = Inches(CONTENT_WIDTH_IN - 0.48)
    number_cell, text_cell = table.rows[0].cells
    set_cell_shading(number_cell, NAVY)
    number_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    np = number_cell.paragraphs[0]
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(np, str(number), bold=True, color=WHITE, size=10)
    set_cell_shading(text_cell, ICE)
    set_cell_margins(text_cell, top=80, bottom=80, start=130, end=130)
    p = text_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    add_text(p, title, bold=True, color=NAVY, size=9.5)
    p2 = text_cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_text(p2, body, color=MUTED, size=8.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(
    doc,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    *,
    font_size=8.2,
    header_fill=NAVY,
    alternate=True,
    alignments: list[int] | None = None,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for index, (cell, header, width) in enumerate(zip(hdr.cells, headers, widths)):
        table.columns[index].width = Inches(width)
        cell.width = Inches(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell, top=80, bottom=80, start=90, end=90)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = alignments[index] if alignments else WD_ALIGN_PARAGRAPH.LEFT
        add_text(p, header, bold=True, color=WHITE, size=8)

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        fill = WHITE if not alternate or row_index % 2 == 0 else ICE
        for index, (cell, value, width) in enumerate(zip(cells, values, widths)):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_shading(cell, fill)
            set_cell_margins(cell, top=75, bottom=75, start=90, end=90)
            set_cell_border(
                cell,
                bottom={"val": "single", "sz": "4", "color": LINE},
            )
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = alignments[index] if alignments else WD_ALIGN_PARAGRAPH.LEFT
            add_text(p, value, color=INK, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_status_row(doc, title: str, percent: str, done: str, missing: str, fill: str, accent: str):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    table.columns[0].width = Inches(1.22)
    table.columns[1].width = Inches(CONTENT_WIDTH_IN - 1.22)
    status_cell, content_cell = table.rows[0].cells
    set_cell_shading(status_cell, accent)
    status_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    sp = status_cell.paragraphs[0]
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(sp, percent, bold=True, color=WHITE, size=18)
    sp2 = status_cell.add_paragraph()
    sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp2.paragraph_format.space_after = Pt(0)
    add_text(sp2, "AVANCE", bold=True, color=WHITE, size=7.2)

    set_cell_shading(content_cell, fill)
    set_cell_margins(content_cell, top=110, bottom=110, start=150, end=150)
    cp = content_cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(4)
    add_text(cp, title, bold=True, color=NAVY, size=11)
    done_p = content_cell.add_paragraph()
    done_p.paragraph_format.space_after = Pt(3)
    add_text(done_p, "Hecho. ", bold=True, color=GREEN, size=8.8)
    add_text(done_p, done, color=INK, size=8.8)
    miss_p = content_cell.add_paragraph()
    miss_p.paragraph_format.space_after = Pt(0)
    add_text(miss_p, "Pendiente. ", bold=True, color=AMBER, size=8.8)
    add_text(miss_p, missing, color=INK, size=8.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.30)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9.7)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(hp, "RADAR DE OPORTUNIDADES  /  INFORME DE AVANCE FASE 1", bold=True, color=NAVY, size=7.8)
    add_text(hp, "     24 JUL 2026", bold=True, color=MID_BLUE, size=7.8)
    p_pr = hp._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), LINE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    table = footer.add_table(rows=1, cols=2, width=Inches(CONTENT_WIDTH_IN))
    set_table_fixed(table)
    table.columns[0].width = Inches(5.9)
    table.columns[1].width = Inches(1.1)
    left, right = table.rows[0].cells
    left.width = Inches(5.9)
    right.width = Inches(1.1)
    lp = left.paragraphs[0]
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "Documento de estado verificable · MVP funcional / piloto controlado", color=MUTED, size=7.2)
    add_page_number(right.paragraphs[0])
    for run in right.paragraphs[0].runs:
        set_run(run, size=7.2, color=MUTED)
    footer._element.remove(fp._element)

    first_footer = section.first_page_footer
    ffp = first_footer.paragraphs[0]
    ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(ffp, "Información verificada al 24 de julio de 2026 · Confidencial para presentación al cliente", color=MUTED, size=7.2)

    core = doc.core_properties
    core.title = "Radar de Oportunidades — Informe de avance Fase 1"
    core.subject = "Trazabilidad, stack, QA, riesgos y hoja de ruta"
    core.author = "Equipo Radar CRECE"
    core.keywords = "Radar CRECE, inmuebles, avance, Fase 1, producción, QA"


def build_cover(doc: Document) -> None:
    brand = doc.add_table(rows=1, cols=2)
    brand.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(brand)
    brand.columns[0].width = Inches(0.35)
    brand.columns[1].width = Inches(CONTENT_WIDTH_IN - 0.35)
    dot, name = brand.rows[0].cells
    set_cell_shading(dot, GOLD)
    set_cell_shading(name, NAVY)
    set_cell_margins(dot, top=70, bottom=70, start=10, end=10)
    set_cell_margins(name, top=70, bottom=70, start=130, end=130)
    p = name.paragraphs[0]
    add_text(p, "RADAR DE OPORTUNIDADES · SISTEMA CRECE", bold=True, color=WHITE, size=9)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(28)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    add_text(p, "INFORME DE AVANCE", bold=True, color=MID_BLUE, size=10)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 0.95
    add_text(title, "Radar de Oportunidades\nInmobiliarias", bold=True, color=NAVY, size=28)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(22)
    add_text(sub, "Trazabilidad de Fase 1, stack técnico, evidencia de QA y plan de endurecimiento", color=MUTED, size=12)

    add_callout(
        doc,
        "Avance real de Fase 1",
        "El MVP supera el objetivo solicitado del 80%. El cálculo excluye explícitamente los elementos aprobados para Fase 2.",
        fill=PALE_BLUE,
        accent=NAVY,
        value="84%",
    )

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(meta)
    meta.columns[0].width = Inches(1.65)
    meta.columns[1].width = Inches(CONTENT_WIDTH_IN - 1.65)
    for row, (label, value) in zip(
        meta.rows,
        [
            ("Preparado para", "Andrés Giraldo"),
            ("Estado", "MVP funcional en producción · apto para demo y piloto controlado"),
            ("Corte", "24 de julio de 2026"),
            ("Producción", "joinsclee-radar.juno8i.easypanel.host"),
        ],
    ):
        left, right = row.cells
        set_cell_shading(left, ICE)
        set_cell_shading(right, WHITE)
        set_cell_margins(left, top=90, bottom=90, start=110, end=110)
        set_cell_margins(right, top=90, bottom=90, start=110, end=110)
        set_cell_border(left, bottom={"val": "single", "sz": "4", "color": LINE})
        set_cell_border(right, bottom={"val": "single", "sz": "4", "color": LINE})
        add_text(left.paragraphs[0], label.upper(), bold=True, color=MID_BLUE, size=7.5)
        add_text(right.paragraphs[0], value, bold=True, color=INK, size=9)

    doc.add_paragraph().paragraph_format.space_after = Pt(9)
    add_callout(
        doc,
        "Lectura honesta",
        "La aplicación está desplegada, usa datos reales y sus recorridos críticos pasan en local y producción. Todavía no debe venderse como un sistema con SLA empresarial ni como 100% endurecido para campañas masivas.",
        fill=PALE_AMBER,
        accent=AMBER,
    )


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    build_cover(doc)

    add_page_break(doc)
    add_section_heading(doc, "01", "Resumen ejecutivo", "Qué puede afirmarse hoy — y qué no")
    add_callout(
        doc,
        "Conclusión para la reunión",
        "Sí hemos superado el 80%: el avance ponderado es 83,6%, comunicado como 84% con un rango razonable de 82–85%.",
        fill=PALE_GREEN,
        accent=GREEN,
        value="84%",
    )
    add_body(
        doc,
        "El Radar ya es un producto demostrable: consulta Supabase en vivo, reúne Portal, activos bancarios y remates, calcula oportunidades, aplica reglas CRECE, permite filtrar, revisar fichas, guardar inmuebles, simular costos e iniciar sesión por correo o Google."
    )
    add_subheading(doc, "Qué sí está listo")
    for text in [
        "Demo completa en el dominio público con datos reales y recorridos coherentes en escritorio y móvil.",
        "Motor estadístico y reglas de negocio verificadas por pruebas automatizadas.",
        "Autenticación por correo y Google OAuth; favoritos anónimos y sincronizables.",
        "Baseline de seguridad HTTP, sanitización de datos y barrera contra valores financieros imposibles.",
        "Suite E2E reproducible que pasa tanto en local como sobre producción.",
    ]:
        add_bullet(doc, text, marker_color=GREEN)
    add_subheading(doc, "Qué no debe prometerse todavía", color=AMBER)
    for text in [
        "SLA empresarial, alta disponibilidad o recuperación probada ante incidentes.",
        "Protección completa ante abuso masivo de login, registro o análisis IA.",
        "Trazabilidad de eventos cerrada al 100% en todos los scrapers y boletines AVAL.",
        "Alertas por correo entregadas y observabilidad centralizada.",
    ]:
        add_bullet(doc, text, marker_color=AMBER)
    add_callout(
        doc,
        "Posicionamiento recomendado",
        "“MVP funcional en producción, apto para demo y piloto controlado; con un plan de endurecimiento concreto para apertura comercial.”",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    add_page_break(doc)
    add_section_heading(doc, "02", "Método de avance", "Porcentaje ponderado sobre la Fase 1 acordada")
    add_body(
        doc,
        "El cálculo no es una suma de pantallas. Cada bloque se pondera por su peso en las HU y por su importancia para operar el MVP. Los elementos aprobados para Fase 2 no se cuentan como deuda."
    )
    score_rows = [
        ["Núcleo CRECE y comparables", "20%", "95%", "19,0"],
        ["Portal, rangos y vigencia", "15%", "78%", "11,7"],
        ["Activos bancarios generales", "15%", "82%", "12,3"],
        ["Grupo AVAL", "10%", "72%", "7,2"],
        ["Remates judiciales", "15%", "92%", "13,8"],
        ["Producto, UX, autenticación y freemium", "15%", "88%", "13,2"],
        ["Producción, seguridad, QA y operación", "10%", "64%", "6,4"],
        ["TOTAL PONDERADO", "100%", "", "83,6%"],
    ]
    add_table(
        doc,
        ["Bloque", "Peso", "Cumplimiento", "Aporte"],
        score_rows,
        [3.85, 0.85, 1.15, 1.15],
        font_size=8.4,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_subheading(doc, "Reglas del cálculo")
    for text in [
        "Completo: existe en código, está conectado al flujo y tiene evidencia de prueba.",
        "Parcial: existe la base técnica, pero falta integración total, automatización u operación.",
        "Pendiente: no está conectado al producto o requiere una decisión/autorización del negocio.",
        "Fuera de alcance: fue ubicado explícitamente en Fase 2 y no reduce el 84%.",
    ]:
        add_bullet(doc, text)
    add_callout(
        doc,
        "Rango defendible",
        "La cifra central es 83,6%. Dependiendo de cuánto peso se asigne al endurecimiento operativo, la evaluación razonable se mueve entre 82% y 85%; por eso 84% es una comunicación equilibrada.",
        fill=PALE_AMBER,
        accent=AMBER,
    )

    add_page_break(doc)
    add_section_heading(doc, "03", "Evidencia objetiva", "Datos, pruebas y controles verificables al corte")
    evidence = [
        ["Portal abierto", "104.626 listados"],
        ["Oportunidades", "18.519 detectadas"],
        ["Oportunidades fuertes", "1.555"],
        ["Activos bancarios", "501"],
        ["Remates judiciales", "688"],
        ["Cobertura API", "67 ciudades"],
        ["Pruebas unitarias/integración", "80/80 aprobadas"],
        ["E2E local", "5/5 aprobadas"],
        ["E2E producción", "5/5 aprobadas"],
        ["TypeScript", "Sin errores"],
        ["Errores de consola E2E", "0"],
        ["Bloqueantes P0/P1 en flujos probados", "0"],
    ]
    add_table(doc, ["Indicador", "Resultado"], evidence, [4.6, 2.4], font_size=9)
    add_subheading(doc, "Recorridos E2E automatizados")
    for text in [
        "Carga del dashboard, endpoint de salud y configuración pública.",
        "Navegación Portal → Bancos → Remates y coherencia de filtros.",
        "Guardado anónimo, persistencia tras recarga y vista Guardados.",
        "Reglas distintas de contraseña para registro/login y disponibilidad de Google.",
        "Navegación y filtros utilizables en viewport móvil de 375 × 812.",
    ]:
        add_bullet(doc, text, marker_color=GREEN)
    add_callout(
        doc,
        "No se alteraron datos de producción",
        "La suite no crea cuentas, no ejecuta scrapers y no escribe en Supabase. La prueba de guardados anónimos usa un contexto limpio y almacenamiento local del navegador.",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    add_page_break(doc)
    add_section_heading(doc, "04", "Trazabilidad de las HU", "Núcleo, Portal y vigencia")
    add_status_row(
        doc,
        "Núcleo del Motor de Comparables",
        "95%",
        "Fórmula CRECE; mediana basada en Portal; exclusión de bancos/remates del universo de referencia; outliers; atributos comparables; cascada; 11 categorías; corte 0,71–0,75 resuelto; persistencia del índice.",
        "Auditar persistencia del 100% de filas históricas y formalizar la decisión del radio/zonas colindantes.",
        PALE_GREEN,
        GREEN,
    )
    add_status_row(
        doc,
        "Rango de precio y vigencia de Portal",
        "78%",
        "Configuración por ciudad en base; niveles 1/2; precios y estratos 2–5; reglas 30/90 días; eventos modelados; cadencia semanal.",
        "Conectar de forma uniforme la escritura de cada evento de ciclo de vida desde los scrapers y verificar activación/población de ciudades.",
        PALE_AMBER,
        AMBER,
    )
    add_subheading(doc, "Decisiones ya resueltas")
    decisions = [
        ["Vacío 0,71–0,75", "Se cerró dentro de Oportunidad Fuerte, con corte inclusivo 0,75."],
        ["Cartagena y Bucaramanga", "Configuradas como Nivel 1 en la migración del 20-jul-2026."],
        ["Cadencia de bancos", "Unificada a 7 días y aplicada en producción."],
        ["Datos bancarios sin estrato", "No bloquean ni excluyen el inmueble."],
    ]
    add_table(doc, ["Decisión", "Tratamiento vigente"], decisions, [2.2, 4.8], font_size=8.8)
    add_callout(
        doc,
        "Punto de trazabilidad",
        "La tabla `oportunidades_historial` y los tipos de evento existen; la deuda es asegurar que todas las fuentes escriban allí de forma homogénea, no diseñar el modelo desde cero.",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    add_page_break(doc)
    add_section_heading(doc, "05", "Trazabilidad de las HU", "Bancos y boletín de Grupo AVAL")
    add_status_row(
        doc,
        "Motor de Activos Bancarios",
        "82%",
        "Fuentes normalizadas; estrato opcional; interfaz sin filtro de estrato; clasificación contra mercado abierto; frescura “Verificado hace…”; rotación y cadencia semanal.",
        "Cerrar identidad secundaria para slugs inestables y uniformar retiro/cambio de precio con historial.",
        PALE_AMBER,
        AMBER,
    )
    add_status_row(
        doc,
        "Grupo AVAL",
        "72%",
        "Resolución dinámica del PDF; extracción por página; código estable; precio, tipo, área, dirección y ciudad; `area_tipo`; deduplicación e imagen por ficha.",
        "Normalización ciudad-departamento; extracción opcional marcada; diff entre boletines; validación de insignias y trazabilidad por versión.",
        PALE_AMBER,
        AMBER,
    )
    add_subheading(doc, "Por qué AVAL no se declara completo")
    add_body(
        doc,
        "El extractor ya produce inventario útil y evita el error crítico de comparar área de terreno contra área construida. Sin embargo, la HU también exige un modelo de snapshot: comparar el boletín completo contra el anterior y registrar NUEVO, VISTO, PRECIO_ACTUALIZADO y RETIRADO. Esa integración aún debe cerrarse."
    )
    add_callout(
        doc,
        "Dato actual",
        "Producción reporta 501 activos bancarios. El volumen demuestra ingesta real, pero no reemplaza la obligación de cerrar la trazabilidad del ciclo de vida.",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    add_page_break(doc)
    add_section_heading(doc, "06", "Trazabilidad de las HU", "Remates, producto y experiencia")
    add_status_row(
        doc,
        "Motor de Remates Judiciales",
        "92%",
        "Origen del demandante, tipo de proceso, cuota-parte, regla conservadora, matriz gratis/pago, alerta visible, barrera de valores imposibles, calculadora y UI responsive.",
        "Cuarentena de anomalías en ingestión/base y reemplazo de imágenes genéricas por evidencia neutral cuando sea posible.",
        PALE_GREEN,
        GREEN,
    )
    add_status_row(
        doc,
        "UX, autenticación y freemium",
        "88%",
        "Navegación responsive; filtros progresivos; foco/teclado; login móvil; favoritos anónimos y sincronización; simulaciones; personalización; correo y Google OAuth.",
        "Entrega real de alertas por correo, onboarding posterior al registro y auditoría WCAG formal.",
        PALE_GREEN,
        GREEN,
    )
    add_subheading(doc, "Mejoras de experiencia aplicadas")
    for text in [
        "El valor y los resultados aparecen antes de exigir una cuenta.",
        "La jerarquía visual prioriza resultados, filtros y acciones.",
        "Los filtros secundarios se revelan progresivamente y muestran estado.",
        "Los estados de carga, vacío, error y guardado entregan feedback inmediato.",
        "Guardados, simulaciones y preferencias sobreviven antes del registro.",
        "El login separa registro de acceso y Google reduce la fricción.",
    ]:
        add_bullet(doc, text)

    add_page_break(doc)
    add_section_heading(doc, "07", "Stack y arquitectura", "Una base simple, auditable y escalable por capas")
    flow = doc.add_table(rows=2, cols=6)
    flow.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(flow)
    flow_widths = [1.15] * 6
    labels = ["FUENTES", "INGESTA", "SUPABASE", "MOTOR CRECE", "API NODE", "WEB UX"]
    details = ["Portal\nBancos\nRemates", "Firecrawl\nPDF\nPlaywright", "PostgreSQL\nAuth\nStorage", "Mediana\nCascada\nReglas", "JSON\nSeguridad\nHealth", "HTML/CSS/JS\nResponsive"]
    for col in range(6):
        for row in range(2):
            cell = flow.cell(row, col)
            cell.width = Inches(flow_widths[col])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=90, bottom=90, start=45, end=45)
            set_cell_shading(cell, NAVY if row == 0 else PALE_BLUE)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text(p, labels[col] if row == 0 else details[col], bold=row == 0, color=WHITE if row == 0 else INK, size=7.5 if row == 0 else 7.7)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    stack_rows = [
        ["Runtime/backend", "Node.js 20+, TypeScript 5.6, ESM"],
        ["Servidor/API", "HTTP nativo de Node, API JSON y estáticos"],
        ["Frontend", "HTML, CSS y JavaScript nativos; responsive"],
        ["Datos", "Supabase PostgreSQL, Auth y Storage"],
        ["Extracción", "Firecrawl, SSR, parsers PDF y Playwright"],
        ["Motor", "Estadística robusta, Índice CRECE y reglas TypeScript"],
        ["Validación", "Zod"],
        ["Automatización", "node-cron + scheduler persistido y cerrojo"],
        ["QA", "node:test, TypeScript, Playwright y Chromium"],
        ["Infraestructura", "Docker, VPS y EasyPanel"],
        ["Entrega", "Git/GitHub, PR y despliegue controlado"],
    ]
    add_table(doc, ["Capa", "Tecnología / decisión"], stack_rows, [2.1, 4.9], font_size=8.7)
    add_callout(
        doc,
        "Criterio arquitectónico",
        "FincaRaíz construye la referencia de mercado; bancos y remates se evalúan contra ella. Esta separación evita contaminar la mediana con ventas forzadas.",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    add_page_break(doc)
    add_section_heading(doc, "08", "Producción, seguridad y QA", "Lo desplegado y los riesgos residuales")
    add_status_row(
        doc,
        "Producción, seguridad y operación",
        "64%",
        "VPS/EasyPanel, Docker, HTTPS, healthcheck, CSP/HSTS, MIME, sanitización, request IDs, errores genéricos, pruebas locales y sobre producción.",
        "Cookies seguras, rate limit, verificación de correo, CI/CD, observabilidad, readiness, no-root, backups, carga, canary y rollback.",
        PALE_AMBER,
        AMBER,
    )
    controls = [
        ["HTTPS / HTTP2", "Verificado", "Operativo en producción"],
        ["CSP + HSTS + anti-framing", "Verificado", "Cabeceras presentes"],
        ["MIME y errores 500", "Verificado", "Pruebas automatizadas"],
        ["Sanitización de datos/URLs", "Implementado", "Mantener pruebas adversariales"],
        ["Sesión HttpOnly", "Pendiente", "Prioridad alta"],
        ["Rate limiting", "Pendiente", "Login, registro y análisis IA"],
        ["CI/CD", "Pendiente", "Suite lista para integrarse"],
        ["Backups/rollback ensayados", "Pendiente", "Requisito para SLA"],
    ]
    add_table(doc, ["Control", "Estado", "Observación"], controls, [2.55, 1.35, 3.1], font_size=8.2)
    add_callout(
        doc,
        "Dependencias",
        "No hay vulnerabilidades críticas ni altas en el audit de producción. Quedan dos moderadas en `node-cron@3 → uuid`; resolverlas requiere migrar y probar `node-cron@4`, no ejecutar un `--force` antes de la demo.",
        fill=PALE_AMBER,
        accent=AMBER,
    )

    add_page_break(doc)
    add_section_heading(doc, "09", "Ruta de 84% a 100%", "Prioridades concretas, no una lista abierta")
    add_subheading(doc, "84% → 90% · cierre inmediato")
    for number, (title, body) in enumerate(
        [
            ("Sesión segura", "Migrar tokens a cookies HttpOnly/Secure/SameSite y añadir protección CSRF."),
            ("Control de abuso", "Rate limiting y cuotas para autenticación y análisis IA."),
            ("Ciclo de vida AVAL", "Diff de boletines y escritura uniforme de historial."),
            ("Activación", "Enviar alertas por correo y completar onboarding registrado."),
            ("Calidad continua", "CI con typecheck, 80 pruebas y los 5 recorridos E2E."),
        ],
        start=1,
    ):
        add_numbered(doc, number, title, body)

    add_subheading(doc, "90% → 100% · endurecimiento")
    hardening = [
        ["Observabilidad", "Métricas, errores, alarmas y presupuesto de latencia."],
        ["Contenedor", "Multi-stage, usuario no-root, readiness y SIGTERM."],
        ["Recuperación", "Backups restaurados en ensayo y runbook de incidentes."],
        ["Validación", "Carga, navegadores/dispositivos, WCAG y revisión legal."],
        ["Entrega", "Canary, rollback probado y aprobación de lanzamiento."],
        ["Calidad de datos", "Cuarentena de anomalías desde ingestión."],
    ]
    add_table(doc, ["Frente", "Resultado esperado"], hardening, [1.8, 5.2], font_size=8.8)
    add_callout(
        doc,
        "Orden recomendado",
        "Las cinco acciones del tramo 84→90 reducen riesgos reales y son más valiosas que añadir nuevas pantallas antes de la presentación comercial.",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    add_page_break(doc)
    add_section_heading(doc, "10", "Alcance, riesgos y exclusiones", "Evitar que Fase 2 se convierta en deuda ficticia")
    add_subheading(doc, "Fuera de Fase 1")
    exclusions = [
        "Scraping y filtro de canon de arriendo.",
        "Comparador de oportunidades estilo Habímetro.",
        "Aplicación móvil nativa.",
        "Banners comerciales Low Ticket / Tradentia.",
        "Página de planes y pasarela de pago.",
        "Nuevas fuentes no acordadas.",
        "Herramientas adicionales distintas de las ya implementadas.",
    ]
    for item in exclusions:
        add_bullet(doc, item, marker_color=MID_BLUE)
    add_subheading(doc, "Riesgos que sí pertenecen al cierre de Fase 1", color=AMBER)
    risks = [
        ["Sesión en localStorage", "Alta", "CSP mitiga XSS, pero no reemplaza HttpOnly."],
        ["Abuso de IA/autenticación", "Alta", "Faltan límites por usuario/IP y concurrencia."],
        ["Historial incompleto", "Media", "Modelo existe; falta escritura uniforme."],
        ["Operación manual", "Media", "Despliegue automático EasyPanel está deshabilitado."],
        ["Dependencias moderadas", "Media", "Migración controlada a node-cron 4."],
        ["Datos de terceros", "Media", "Los volúmenes y formatos cambian por fuente."],
    ]
    add_table(doc, ["Riesgo", "Nivel", "Tratamiento"], risks, [2.25, 0.8, 3.95], font_size=8.4)
    add_callout(
        doc,
        "Mensaje de alcance",
        "No se está vendiendo un producto terminado al 100%; se está entregando un MVP real al 84%, con brechas identificadas, priorizadas y medibles.",
        fill=PALE_AMBER,
        accent=AMBER,
    )

    add_page_break(doc)
    add_section_heading(doc, "11", "Guion para la demostración", "Una reunión de 12–15 minutos basada en evidencia")
    demo_steps = [
        ("Cifras en vivo", "Abrir producción y mostrar Portal, Bancos y Remates."),
        ("Oportunidad", "Filtrar Portal por ciudad/tipo y explicar descuento/comparables."),
        ("Regla bancaria", "Mostrar que Bancos no excluye por estrato."),
        ("Riesgo judicial", "Abrir Remates y explicar postura, avalúo y cuota-parte."),
        ("Valor antes del registro", "Guardar un inmueble sin cuenta y recargar."),
        ("Persistencia", "Entrar a Guardados y mostrar que el inmueble sigue allí."),
        ("Activación", "Mostrar personalización del Radar y simulador."),
        ("Acceso", "Abrir Login y enseñar correo + Google."),
        ("Cierre ejecutivo", "Presentar el 84% y las cinco acciones para llegar al 90%."),
    ]
    for index, (title, body) in enumerate(demo_steps, start=1):
        add_numbered(doc, index, title, body)
    add_callout(
        doc,
        "Frase de cierre sugerida",
        "“El producto ya encuentra y presenta oportunidades reales. El siguiente tramo no es rehacerlo: es endurecer sesión, operación y trazabilidad para escalar con seguridad.”",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    add_page_break(doc)
    add_section_heading(doc, "12", "Dictamen y fuentes", "Conclusión verificable")
    add_callout(
        doc,
        "DICTAMEN",
        "El proyecto sí ha alcanzado más del 80% de la Fase 1. La evidencia permite defender un 84% real, sin contar funcionalidades de Fase 2 como deuda. Está listo para presentación y piloto controlado.",
        fill=PALE_GREEN,
        accent=GREEN,
        value="84%",
    )
    add_subheading(doc, "Fuentes contractuales y técnicas revisadas")
    sources = [
        "Radar de Oportunidades — Documentación y Propuesta v1.0.",
        "Especificación Técnica del Motor de Comparables v1.2.",
        "Índice Maestro de HU Fase 1.",
        "HU Núcleo del Motor de Comparables.",
        "HU Rango de Precio y Vigencia del Motor Radar.",
        "HU Motor de Activos Bancarios.",
        "HU Motor de Activos Bancarios — Grupo AVAL.",
        "HU Motor de Remates Judiciales.",
        "Código, migraciones y pruebas del repositorio radar-mvp.",
        "APIs y cabeceras del despliegue público verificadas el 24-jul-2026.",
    ]
    for source in sources:
        add_bullet(doc, source, marker_color=MID_BLUE, size=9)
    add_subheading(doc, "Trazas de entrega")
    traces = [
        ["Producción", "https://joinsclee-radar.juno8i.easypanel.host/"],
        ["Versión desplegada", "38d3a878b52504b23cfae5238e621d67e318420d"],
        ["Pruebas", "80 unitarias/integración + 5 E2E local + 5 E2E producción"],
        ["Informe QA", ".gstack/qa-reports/qa-report-joinsclee-radar-2026-07-24.md"],
        ["Informe fuente", "docs/INFORME_AVANCE_FASE1_2026-07-24.md"],
    ]
    add_table(doc, ["Elemento", "Referencia"], traces, [1.65, 5.35], font_size=8.6)
    add_body(
        doc,
        "Nota: los volúmenes son una fotografía del corte y cambian cuando las fuentes se actualizan. El porcentaje se recalculará cuando se cierren los pendientes del tramo 84→90.",
        color=MUTED,
        size=8.5,
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
